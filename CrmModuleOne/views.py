from django.shortcuts import render, get_object_or_404, redirect
from .models import InformationPost
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator
from django.http import JsonResponse,HttpResponse
from django.views.decorators.csrf import csrf_exempt
from .forms import ClientForm, TaskForm , MeetingForm, NoteForm, ClientFileForm, PaymentForm, DocumentTemplateForm, PreleadForm
from docx import Document
from .models import Client, ProductCategory, Product, ProductAttribute,SubsidyProgram, SubsidyOption, ProductConfiguration, SubsidyProductCriteria, ClientFile, Task, Meeting, Payment, Note, DocumentTemplate, Offer, OfferProduct, Profile, DocumentTemplate, ProductDocumentRequirement, ActivityLog, Notification, Prelead, Parcel 
from django.shortcuts import render, get_object_or_404, redirect
from .forms import ProductConfigurationForm
import re,os, docx, locale, pdfkit, json
from decimal import Decimal
from datetime import datetime, timedelta, date
from django.urls import reverse
from django.contrib import messages
from django.core.files import File
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.utils import timezone
from collections import defaultdict
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from io import BytesIO
from django.conf import settings
from django import forms
from .models import DocumentTemplate
from num2words import num2words
import uuid  # ✅ Import dla generowania UUID
import traceback  # ✅ Import dla obsługi błędów
from django.contrib.auth.models import User, Group
import requests
from .services import get_all_rows
import pytz
from django.views.decorators.http import require_http_methods
from django.views.decorators.http import require_POST
from django.shortcuts import redirect
from django.contrib import messages
from CrmModuleOne.models import Parcel
from CrmModuleOne.utils import get_coordinates
import time

@login_required
def prelead_form(request, lead_id):
    lead = get_object_or_404(Prelead, pk=lead_id)
    context = {
        'lead': lead,
        'handlowcy': User.objects.filter(groups__name="Handlowcy"),
    }
    return render(request, 'prelead_form.html', context)

@login_required
def grafiks(request):
    handlowcy = Group.objects.get(name="handlowcy").user_set.all().order_by("username")
    leads = Prelead.objects.select_related('parcel').all()

    parcels_data = []
    for p in Parcel.objects.filter(latitude__isnull=False, longitude__isnull=False).select_related('lead'):
        lead_status = None
        if p.lead_id:
            lead_status = p.lead.status  # jeśli powiązanie istnieje
    
        parcels_data.append({
            "latitude": p.latitude,
            "longitude": p.longitude,
            "town": p.town,
            "plot_number": p.plot_number,
            "area": float(p.area),
            "status": lead_status or "nieokreślony"
        })

    context = {
        'leads': leads,
        'handlowcy': handlowcy,
        'parcels_json': parcels_data
    }

    print("📦 Dane działek:", json.dumps(parcels_data, indent=2, ensure_ascii=False))
    return render(request, 'grafiks.html', context)



@login_required
def get_prelead_detail(request, lead_id):
    try:
        lead = Prelead.objects.get(id=lead_id)
        parcels = Parcel.objects.filter(lead=lead)

        return JsonResponse({
            "success": True,
            "lead": {
                "id": lead.id,
                "first_name": lead.first_name,
                "last_name": lead.last_name,
                "phone": lead.phone,
                "email": lead.email,
                "status": lead.status,
                "potential": lead.potential
            },
            "parcels": list(parcels.values("plot_number", "area", "town"))
        })

    except Prelead.DoesNotExist:
        return JsonResponse({"success": False, "error": "Lead nie istnieje"})

@login_required
def list_parcels(request):
    lead_id = request.GET.get('prelead_id')
    try:
        lead = Prelead.objects.get(id=lead_id)
    except Prelead.DoesNotExist:
        return JsonResponse({'success': False, 'error': f'Lead nie istnieje{lead_id}'})

    parcels = lead.parcels.all().values(
        'id', 'voivodeship', 'county', 'town', 'precinct', 'plot_number', 'area'
    )
    return JsonResponse({'success': True, 'parcels': list(parcels)})



@require_POST
@login_required
def delete_parcel(request):
    try:
        parcel_id = request.POST.get('parcel_id')
        if not parcel_id:
            return JsonResponse({'success': False, 'error': 'Brak ID działki'})

        Parcel.objects.get(id=parcel_id).delete()
        return JsonResponse({'success': True})
    except Parcel.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Działka nie istnieje'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


@require_POST
@login_required
def add_parcel(request):
    if request.method == 'POST':
        try:
            lead_id = request.POST.get('prelead_id')
            try:
                lead = Prelead.objects.get(id=lead_id)
            except Prelead.DoesNotExist:
                return JsonResponse({'success': False, 'error': f'Lead nie istnieje{lead_id}'})

            Parcel.objects.create(
                lead=lead,
                voivodeship=request.POST.get('voivodeship_label'),
                county=request.POST.get('county_label'),
                town=request.POST.get('commune_label'),
                precinct=request.POST.get('precinct'),
                plot_number=request.POST.get('plot_number'),
                area=request.POST.get('area')
            )

            return JsonResponse({'success': True})
        except Prelead.DoesNotExist:
            return JsonResponse({'success': False, 'error': f'Lead nie istnieje{lead_id}'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})





@login_required
@require_POST
def save_prelead(request):
    try:
        prelead_id = request.POST.get('prelead_id')
        if not prelead_id:
            return JsonResponse({'success': False, 'error': 'Brak ID preleada'}, status=400)

        try:
            lead = Prelead.objects.get(id=prelead_id)
        except Prelead.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Lead nie istnieje'}, status=404)

        # Uprawnienia
        if not request.user.is_staff and lead.user != request.user:
            return JsonResponse({'success': False, 'error': 'Brak uprawnień'}, status=403)

        # Aktualizacja podstawowych pól
        lead.potential = request.POST.get('potential', lead.potential)
        lead.note = request.POST.get('note', lead.note)

        note_prefix = ""
        log_updated = False

        # Obsługa statusu
        status_action = request.POST.get('status_update')
        if status_action:
            if lead.status in ['NUM', 'UM'] and status_action in ['no_answer', 'inactive_number']:
                pass  # Nie zmieniaj statusu
            elif status_action == 'no_answer':
                if lead.status == 'NO1':
                    lead.status = 'NO2'
                elif lead.status == 'NO2':
                    lead.status = 'NO3'
                else:
                    lead.status = 'NO1'
                note_prefix = "Nie odbiera"
                log_updated = True
            elif status_action == 'inactive_number':
                lead.status = 'NN2' if lead.status == 'NN' else 'NN'
                note_prefix = "Numer nieaktywny"
                log_updated = True
            elif status_action == 'to_verify':
                lead.status = 'NUM'
                note_prefix = "Do weryfikacji"
                log_updated = True
            elif status_action == 'contract_sent':
                lead.status = 'UM'
                note_prefix = "Umowa wysłana"
                log_updated = True

        # Dodanie loga jeśli status się zmienił
        if log_updated and note_prefix:
            timestamp = timezone.now().strftime("%Y-%m-%d %H:%M")
            user_info = request.user.get_full_name() or request.user.username
            log_entry = f"[{timestamp}] {user_info}: {note_prefix}"
            lead.log = f"{log_entry}\n{lead.log or ''}"

        # Przypisanie użytkownika
        assigned_to_id = request.POST.get('assigned_to')
        client_created = False

        if assigned_to_id:
            try:
                assigned_user = User.objects.get(id=assigned_to_id)
                lead.user = assigned_user
            except User.DoesNotExist:
                return JsonResponse({'success': False, 'error': 'Nieprawidłowy użytkownik'}, status=400)

            # Stworzenie klienta
            client = Client.objects.create(
                first_name=lead.first_name,
                last_name=lead.last_name or "",
                phone=lead.phone,
                city=lead.city,
                email=lead.email,
                user=assigned_user,
                status='lead',
                modified_by=request.user,
                potential=lead.potential,
            )
            # Tworzenie notatki z danymi z preleada
            if lead.note:
                Note.objects.create(
                    client=client,
                    text=lead.note,
                    author=request.user,
                    is_important=False  # lub True, jeśli chcesz oznaczyć je jako ważne
                )
            lead.status = 'PR'
            timestamp = timezone.now().strftime("%Y-%m-%d %H:%M")
            user_info = request.user.get_full_name() or request.user.username
            log_entry = f"[{timestamp}] {user_info}: Przekształcono na klienta"
            lead.log = f"{log_entry}\n{lead.log or ''}"
            client_created = True

        # Zapisz zmiany
        lead.save()

        if client_created:
            messages.success(request, "Prelead został przypisany jako klient.")

        # Przygotuj log do odpowiedzi
        full_log = lead.log or ''
        logs = [log.strip() for log in full_log.split('\n') if log.strip()]
        last_log = logs[-1] if logs else 'Brak logów'

        show_modal = request.session.pop('show_import_modal', False)

        return JsonResponse({
            'success': True,
            'log': full_log,
            'last_log': last_log,
            'potential': lead.potential,
            'status': lead.status,
            'assigned_to': lead.user.id if lead.user else None,
            'assigned_to_name': lead.user.get_full_name() if lead.user else None,
            'show_import_modal': show_modal
        })

    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@require_http_methods(["POST"])
def import_leadsfb(request):
    leads = get_all_rows("LeadyFB")
    new_leads = 0
    duplicates = 0
    errors = 0

    for lead in leads:
        try:
            digits = re.sub(r'\D', '', lead["phone"])
            phone = digits[-9:] if len(digits) >= 9 else None

            parts = lead["fullname"].strip().split()
            first_name = parts[0]
            last_name = parts[1] if len(parts) > 1 else ''
            
            created_time = datetime.fromisoformat(lead["created_time"]).astimezone(pytz.UTC)

            if not phone or len(phone) != 9:
                errors += 1
                continue

            if Prelead.objects.filter(phone=phone).exists():
                duplicates += 1
                continue

            Prelead.objects.create(
                first_name=first_name,
                last_name=last_name,
                phone=phone,
                email = lead["email"],
                created_at=created_time,
                potential='medium',
                status="ST",
            )
            new_leads += 1

        except Exception as e:
            errors += 1
            print(f"Błąd: {str(e)}")

    result_message = (
        f"Zaimportowano {new_leads} nowych leadów. "
        f"Duplikaty: {duplicates}. "
        f"Błędów: {errors}."
    )

    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        return JsonResponse({"success": True, "message": result_message})
    else:
        messages.success(request, result_message)
        return redirect('leadsfb')

@login_required
def leadsfb(request):
    tab = request.GET.get('tab', 'leady')
    subtab = request.GET.get('subtab')

    handlowcy = Group.objects.get(name="handlowcy").user_set.all().order_by("username")

    leads = Prelead.objects.all()
    if tab == 'archiwalne':
        if subtab == 'no_answer':
            leads = leads.filter(status__in=['NO3'])
        elif subtab == 'inactive_number':
            leads = leads.filter(status__in=['NN2'])
        elif subtab == 'przypisane':
            leads = leads.filter(status__in=['PR'])
        else:
            leads = leads.filter(status__in=['NO3','NN2','PR'])
    elif tab == 'weryfikacja':
        leads = leads.filter(status='NUM')
    elif tab == 'analiza_umowy':
        leads = leads.filter(status='UM')
    else:
        leads = leads.filter(status__in=['NO1', 'NO2', 'NN', 'ST'])

    paginator = Paginator(leads.order_by('-created_at'), 10)  # Paginacja leadów
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    context = {
        'leads': page_obj,
        'handlowcy': handlowcy,
        'tab': tab,
        'subtab': subtab,
    }
    return render(request, 'leadsfb.html', context)


def kwota_slownie(kwota):
    """
    Konwertuje liczbę na jej słowną reprezentację w języku polskim.
    
    :param kwota: Liczba do konwersji (może być int lub float)
    :return: Słowna reprezentacja kwoty
    """
    # Podział na złote i grosze
    zlote = int(kwota)
    grosze = round((kwota - zlote) * 100)

    # Konwersja liczb na tekst
    zlote_slownie = num2words(zlote, lang='pl')
    grosze_slownie = num2words(grosze, lang='pl') if grosze > 0 else "zero"

    # Dostosowanie formatu końcowego
    if zlote == 1:
        zlote_sufix = "złoty"
    elif 2 <= zlote % 10 <= 4 and (zlote < 10 or zlote > 20):
        zlote_sufix = "złote"
    else:
        zlote_sufix = "złotych"

    if grosze == 1:
        grosze_sufix = "grosz"
    elif 2 <= grosze % 10 <= 4 and (grosze < 10 or grosze > 20):
        grosze_sufix = "grosze"
    else:
        grosze_sufix = "groszy"

    return f"{zlote_slownie} {zlote_sufix} {grosze_slownie} {grosze_sufix}"


class MissingDataForm(forms.Form):
    """Dynamiczny formularz do uzupełniania brakujących danych."""
    
    def __init__(self, missing_fields, *args, **kwargs):
        super().__init__(*args, **kwargs)
        for field in missing_fields:
            self.fields[field] = forms.CharField(
                required=True,
                label=field.replace("_", " ").capitalize(),
                widget=forms.TextInput(attrs={'class': 'form-control'})
            )

@login_required
@csrf_exempt
def generate_offer_pdf(request):
    if request.method == 'POST':
        try:
            # **1️⃣ Odbierz dane z żądania**
            data = json.loads(request.body)
            lead_data = data.get('customer', {}).get('leadData')  # Istniejący lead
            form_data = data.get('customer', {})  # Dane z formularza
 # Pobieranie dynamicznych pól
            required_fields = form_data.get("required_fields", {})

            print(f"📌 Otrzymane wymagane pola: {required_fields}")  # Debugowanie w konsoli

            print(f"🔹 Odebrane dane: {data}")

            # **2️⃣ Pobranie lub wygenerowanie unique_id**
            unique_id = lead_data.get('unique_id') if lead_data else form_data.get('unique_id', str(uuid.uuid4()))
            print(f"🔹 Unique ID klienta: {unique_id}")

            # Pobranie klienta lub stworzenie nowego, ustawienie domyślnego użytkownika
            client, created = Client.objects.get_or_create(
                unique_id=unique_id,
                defaults={
                    "first_name": form_data.get("first_name", ""),
                    "last_name": form_data.get("last_name", ""),
                    "phone": form_data.get("phone", ""),
                    "email": form_data.get("email", ""),
                    "street": form_data.get("street", ""),
                    "house_number": form_data.get("house_number", ""),
                    "city": form_data.get("city", ""),
                    "postal_code": form_data.get("postal_code", ""),
                    "current_heating_source": form_data.get("current_heating_source") or None,
                    "construction_permission_year": form_data.get("construction_permission_year") or None,
                    "land_registry_number": form_data.get("land_registry_number") or None,
                    "plot_number": form_data.get("plot_number") or None,
                    "income_threshold": form_data.get("income_threshold", ""),
                    "household_members": form_data.get("household_members") or None,
                    "farm_conversion_hectares": form_data.get("farm_conversion_hectares") or None,
                    "income_per_person": form_data.get("income_per_person") or None,
                    "marital_status": form_data.get("marital_status") or None,
                    "is_sole_owner": form_data.get("is_sole_owner", False),
                    "runs_business": form_data.get("runs_business", False),
                    "has_joint_property": form_data.get("has_joint_property", False),
                    "potential": form_data.get("potential", "low"),
                    "person_type": form_data.get("person_type", "individual"),
                    "has_different_investment_address": form_data.get("has_different_investment_address", False),
                    "investment_street": form_data.get("investment_street") or None,
                    "investment_house_number": form_data.get("investment_house_number") or None,
                    "investment_postal_code": form_data.get("investment_postal_code") or None,
                    "status": form_data.get("status_record", "lead"),
                    "user": request.user,  # ✅ Dodanie użytkownika
                }
            )

            updated = False
            update_fields = {}

            for key, value in form_data.items():
                if hasattr(client, key) and value not in [None, ""]:
                    current_value = getattr(client, key)
                    if current_value != value:
                        update_fields[key] = value
                        updated = True

            if updated:
                for key, value in update_fields.items():
                    setattr(client, key, value)
                client.save()
                print(f"✅ Klient {client.first_name} {client.last_name} został zaktualizowany.")

            # **✅ Przeniesienie `if created:` poza `if not created:`**
            if created:
                print(f"✅ Utworzono nowego klienta: {client.first_name} {client.last_name}")
            
            # Oblicz bieżącą datę i datę ważności
            locale.setlocale(locale.LC_TIME, "pl_PL.UTF-8")
            current_date = datetime.now().strftime("%d.%m.%Y %H:%M")
            valid_until_date = (datetime.now() + timedelta(days=8)).strftime("%d %B %Y")


            # Pobranie listy produktów z oferty klienta
            products = form_data.get("subsidy_results", {}).get("subsidy_results", {}).get("subsidy_distribution", [])

            # Tworzenie nowej oferty
            offer = Offer.objects.create(
                client=client,
                total_price=data["customer"]["subsidy_results"]["subsidy_results"]["total_gross_price"],
                must_payment=data["customer"]["subsidy_results"]["subsidy_results"]["mustPayment"],
                total_margin = data["customer"]["total_margin"],
                additional_terms= data["customer"].get("addInfo", "")
            )
            
            postal_code = data.get("required_fields", {}).get("postal-code", None) or data.get("customer", {}).get("postal_code", None)

                        # Ustalenie tytułu na podstawie częściowej zgodności nazw
            has_pv = any("Instalacja PV" in product["product_name"] for product in products)
            has_storage = any("Magazyn Energii" in product["product_name"] for product in products)
            if has_pv and has_storage:
                subtitle = "FOTOWOLTAICZNEGO Z MAGAZYNEM ENERGII"
            elif has_pv:
                subtitle = "FOTOWOLTAICZNEGO"
            elif has_storage:
                subtitle = "MAGAZYNU ENERGII"
            else:
                subtitle = "BRAK PRODUKTÓW W OFERCIE"
            if has_pv or has_storage:
                                # Pobranie danych o nasłonecznieniu
                solar_data = get_solar_data(data["customer"]["postal_code"])
                if solar_data is None:
                    print(f"⚠️ Błąd pobierania danych o nasłonecznieniu dla kodu: {data['customer']['postal_code']}")

                lat, lon, annual_solar_radiation = solar_data
                print(f"📍 Współrzędne: {lat}, {lon} | ☀️ Nasłonecznienie roczne: {annual_solar_radiation} kWh/m²")
                for product in products:
                    if "Moc instalacji" in product["attributes"]:
                        moc_instalacji_str = product["attributes"]["Moc instalacji"]



                        # Parsowanie "Moc instalacji" - obsługa różnych myślników
                        match = re.search(r'(\d+)\s*szt\.\s*[–-]\s*([\d,.]+)\s*kWp', moc_instalacji_str)

                        if match:
                            try:
                                panel_value = int(match.group(1))  # Ilość paneli
                                foto_power = float(match.group(2).replace(",", "."))  # Moc instalacji
                                print(f"✅ Ilość paneli: {panel_value}, Moc instalacji: {foto_power} kWp")
                            except ValueError:
                                print(f"⚠️ Błąd konwersji '{match.group(2)}' na float!")
                                panel_value = None
                                foto_power = None
                        else:
                            print(f"⚠️ Nie udało się sparsować 'Moc instalacji': {moc_instalacji_str}")
                            panel_value = None  # Zapobieganie błędowi "UnboundLocalError"
                            foto_power = None

                        # Pobranie falownika i paneli (z obsługą brakujących danych)
                        inverter_name = product["attributes"].get("Falownik", "Nieznany falownik")
                        panels_name = product["attributes"].get("Panele", "Nieznane panele")
                    else:
                        foto_power = 0
                    if "Moc magazynu" in product["attributes"]:
                        mag_name = product["attributes"].get("Model magazynu", "Nieznane panele")

                
                elec_bill = float(required_fields["elec_bill"])
                kWhPrice  = float(required_fields["kWhPrice"])
                try:
                    pv_exist = float(required_fields.get("pvExist", "0").replace(",", "."))
                except ValueError:
                    pv_exist = 0.0

                foto_power = float(foto_power) + pv_exist
                estimate_electric_raise_year = float(required_fields["estimate_electric_raise_year"])/100
                postal_code = required_fields["postal-code"]
                print(elec_bill, (kWhPrice), (estimate_electric_raise_year), (postal_code))
                y_axis_values = [
                    round((foto_power * annual_solar_radiation * 0.85) / 12),  # Zaokrąglenie wartości
                    round(((foto_power * annual_solar_radiation * 0.85)/12) * 0.8),
                    round(((foto_power * annual_solar_radiation * 0.85)/12) * 0.6),
                    round(((foto_power * annual_solar_radiation * 0.85)/12) * 0.4),
                    round(((foto_power * annual_solar_radiation * 0.85)/12) * 0.2),
                    0
                ]
                y_axis_positions = [(i * 25) for i in range(len(y_axis_values))]
                estimates = {
                    "1_year": round((foto_power * annual_solar_radiation * 0.85)),
                    "5_years": round((foto_power * annual_solar_radiation * 0.85) * 5),
                    "10_years": round((foto_power * annual_solar_radiation * 0.85) * 10),
                    "20_years": round((foto_power * annual_solar_radiation * 0.85) * 20),}
                km_per_kWh = 5.56  # średni zasięg auta elektrycznego na 1 kWh
                fuel_saving_per_km = 0.08  # oszczędność paliwa na 1 km 
                driving_range = {key: round(value * km_per_kWh) for key, value in estimates.items()}
                fuel_savings = {key: round(driving_range[key] * fuel_saving_per_km) for key in estimates}
                # Roczna redukcja emisji CO2 na podstawie produkcji energii (przyjmujemy 0.4 kg CO2 na 1 kWh)
                co2_reduction_per_kWh = 0.4  
                co2_emission_saved = {key: round(value * co2_reduction_per_kWh) for key, value in estimates.items()}

                # NOx i SOx (przyjmujemy, że 1 kWh zmniejsza NOx i SOx o 5.56 g)
                nox_sox_reduction_per_kWh = 0.00556  
                nox_sox_saved = {key: round(value * nox_sox_reduction_per_kWh) for key, value in estimates.items()}

                # Oszczędność paliwa na podstawie przelicznika 1 litr paliwa = 12.5 kg CO2
                fuel_savings_trees = {key: round(value * 0.08) for key, value in driving_range.items()}
                # Wzrost ceny prądu o 8% rocznie
                growth_rate = 1 + estimate_electric_raise_year  # 1.08 (8% wzrostu rocznie)

                # Roczne oszczędności na początku
                initial_annual_savings = round(((foto_power * annual_solar_radiation * 0.85) * kWhPrice*0.7), 2)  
                # Sumujemy oszczędności przez 20 lat, uwzględniając wzrost cen
                total_savings_20_years = sum([initial_annual_savings * (growth_rate ** year) for year in range(1, 21)])
                total_savings_20_years = round(total_savings_20_years, 2)
                #Suma otrzymanych dopłat do wyprodukowanejenergii w ciągu20lat
                sum_elec_sell = total_savings_20_years*0.35*0.3
                # Suma oszczędności wraz z dotacjami

                total_savings_with_subsidies = round(total_savings_20_years + sum_elec_sell, 2)

                # Koszt energii bez fotowoltaiki (przy wzroście cen)
                energy_cost_without_pv = round( total_savings_20_years, 2)

                # Koszt energii z fotowoltaiką (czyli rachunek po odjęciu rosnących oszczędności)
                energy_cost_with_pv = round(energy_cost_without_pv*0.12, 2)
                
                # Koszt systemu PV po uwzględnieniu dotacji
                try:
                    total_gross_price = float(data["customer"]["subsidy_results"]["subsidy_results"]["total_gross_price"])
                    total_subsidy = float(data["customer"]["subsidy_results"]["subsidy_results"]["total_subsidy"])
                    total_pv_cost = round(total_gross_price - total_subsidy, 2)
                except ValueError:
                    print("Błąd: Nie można przekonwertować wartości na liczbę.")
                    total_pv_cost = None
                cash_flow_data = []
                cash_flow_data_list = []
                cumulative_cash_flow = -total_pv_cost
                for year in range(1, 20 + 1):
                    degradation_factor = (1 - (0.005 * year))  # Degradacja paneli (0.5% rocznie)
                    production = round(foto_power * annual_solar_radiation * 0.85 * degradation_factor, 2)
                    
                    current_kWhPrice = round(kWhPrice * ((1 + estimate_electric_raise_year) ** (year - 1)), 2)
                    savings = round(production * current_kWhPrice, 2)

                    # Dopłata tylko w 1. roku
                    subsidy = round(data["customer"]["subsidy_results"]["subsidy_results"]["total_subsidy"], 2) if year == 1 else 0
                    
                    # Skumulowany przepływ
                    cumulative_cash_flow += savings
                    cash_flow_data_list.append(cumulative_cash_flow)
                    min_value = min(cash_flow_data_list)
                    max_value = max(cash_flow_data_list)
                    # Koszt energii bez PV
                    bill_without_pv = round(round(elec_bill / kWhPrice, 2) * ((1 + estimate_electric_raise_year) ** (year - 1)) * kWhPrice, 2)

                    # Koszt energii z PV - przyjmujemy, że PV pokrywa 85% zużycia
                    bill_with_pv = round(bill_without_pv * 0.03, 2)

                    cash_flow_data.append({
                        "year": year,
                        "production": f"{production:,.2f}",
                        "savings": f"{savings:,.2f}",
                        "subsidy": f"{subsidy:,.2f}",
                        "cumulative_cash_flow": f"{cumulative_cash_flow:,.2f}",
                        "bill_without_pv": f"{bill_without_pv:,.2f}",
                        "bill_with_pv": f"{bill_with_pv:,.2f}"
                    })
                graph_axis_y = [
                    round(min_value),
                    round(min_value * 0.5266),
                    0,
                    round(max_value * 0.298),
                    round(max_value * 0.564),
                    round(max_value * 0.73),
                    round(max_value * 0.90)
                ]
                num_bars = 20
                graph_values = []
                for i in range(num_bars):
                    index = int(i / (num_bars / len(graph_axis_y)))  # Przypisanie indeksu do graph_axis_y
                    proportion = (i % (num_bars / len(graph_axis_y))) / (num_bars / len(graph_axis_y))
                    value = round(graph_axis_y[index] + proportion * (graph_axis_y[min(index+1, len(graph_axis_y)-1)] - graph_axis_y[index]))
                    graph_values.append(value)

                print(graph_values)
                # Oszczędności po odjęciu kosztów systemu
                net_savings = round(total_savings_20_years - total_pv_cost, 2)
                if energy_cost_without_pv == 0:
                    bardDougreen = 0
                else:
                    bardDougreen = round((energy_cost_with_pv / energy_cost_without_pv) * 100)
                # Średnie roczne oszczędności (średnia wartość z 20 lat)
                annual_savings = round(total_savings_20_years / 20, 2)
            # Dane do szablonu
            offer_data = {
                "customer": {
                    "first_name": data["customer"]["first_name"],
                    "last_name": data["customer"]["last_name"],
                    "phone": data["customer"]["phone"],
                    "email": data["customer"]["email"],
                    "street": data["customer"]["street"],
                    "house_number": data["customer"]["house_number"],
                    "city": data["customer"]["city"],
                    "postal_code": postal_code if postal_code else data["customer"]["postal_code"],
                    "paymentInfo": data["customer"].get("paymentInfo", ""),
                    "paymentMethod": data["customer"].get("paymentMethod", ""),
                    "addInfo": data["customer"].get("addInfo", "")
                },
                "logo_url": request.build_absolute_uri('/static/images/HSZ_logo_bez.png'),
                "subsidy_results": data["customer"]["subsidy_results"]["subsidy_results"],
                "total_subsidy": round(Decimal(data["customer"]["subsidy_results"]["subsidy_results"]["total_subsidy"]), 2),
                "total_payment": round(data["customer"]["subsidy_results"]["subsidy_results"]["mustPayment"],2),
                "total_gross_price": round(Decimal(data["customer"]["subsidy_results"]["subsidy_results"]["total_gross_price"]), 2),
                "total_price": round(data["customer"]["subsidy_results"]["subsidy_results"]["total_price"],2),
                "totalVat": round((data["customer"]["subsidy_results"]["subsidy_results"]["total_gross_price"])-(data["customer"]["subsidy_results"]["subsidy_results"]["total_price"]),2),
                "date": current_date,
                "valid_until": valid_until_date,
                
            }
                        # Przekazanie do szablonu
            offer_data["subtitle"] = subtitle
            if has_pv:
                offer_data.update({
                "inverter_name": inverter_name,
                "panels_name":panels_name,
                "panel_value": panel_value,
                "panel_m2":panel_value *1.87,
            })
            if has_storage:
                offer_data.update({
                "mag_name":mag_name,
            })
            if has_pv or has_storage:
                try:
                    denominator1 = elec_bill / kWhPrice
                    field1_value = (foto_power * annual_solar_radiation * 0.85) / denominator1 * 100
                    field1 = 100 if round(field1_value) > 100 else round(field1_value)
                except ZeroDivisionError:
                    field1 = 0

                try:
                    denominator2 = foto_power * annual_solar_radiation * 0.85
                    field2_value = (elec_bill / kWhPrice) / denominator2 * 100
                    field2 = 100 if round(field2_value) > 100 else round(field2_value)
                except ZeroDivisionError:
                    field2 = 0
                offer_data.update({
                "electric_bill": elec_bill,
                "electric_bill_monthly": elec_bill/12,
                "kWh_price": kWhPrice,
                "foto_power": foto_power,
                "foto_power_recom": foto_power + 0.05,
                "estimate_year_production": round(foto_power * annual_solar_radiation * 0.85, 2),
                "estimate_year_consume": round(elec_bill / kWhPrice, 2),
                "solar_radiation": annual_solar_radiation,
                "electric_raise":round(estimate_electric_raise_year*100,1),
                "rasie20Y": round(kWhPrice*(1+estimate_electric_raise_year) ** 20,2),
                "CO_emision": round((elec_bill / kWhPrice)*0.8),
                "night_electric_without_mag": round(((elec_bill / kWhPrice) -((elec_bill / kWhPrice)*0.30))*kWhPrice),
                "night_electric_with_mag": round(((elec_bill / kWhPrice) -((elec_bill / kWhPrice)*0.77))*kWhPrice),
                "saving": round((((elec_bill / kWhPrice) -((elec_bill / kWhPrice)*0.3))*kWhPrice)-(((elec_bill / kWhPrice) -((elec_bill / kWhPrice)*0.77))*kWhPrice)),
                "y_axis_sun": zip(y_axis_values, y_axis_positions),
                "estimates": estimates,
                "driving_range": driving_range,
                "fuel_savings": fuel_savings,
                "co2_emission_saved": co2_emission_saved,
                "nox_sox_saved": nox_sox_saved,
                "fuel_savings": fuel_savings_trees,
                "annual_savings": round(((foto_power * annual_solar_radiation * 0.85)*kWhPrice),2),  # Średnie roczne oszczędności
                "annual_savings": annual_savings,
                "total_savings_20_years": total_savings_20_years,
                "subsidy_total": sum_elec_sell,
                "total_savings_with_subsidies": total_savings_with_subsidies,
                "energy_cost_with_pv": energy_cost_with_pv,
                "energy_cost_without_pv": energy_cost_without_pv,
                "bardDougreen": bardDougreen,
                "net_savings": net_savings,
                "total_pv_cost": total_pv_cost,
                "cash_flow_data": cash_flow_data,
                "graph_axis_y":graph_axis_y,
                "lat":lat,
                "lon":lon,
                "field1": field1,
                "field2": field2,
            })
                


                # Renderowanie HTML
                html_content = render_to_string('offer_template.html', offer_data)
            else: 
                                # Renderowanie HTML
                html_content = render_to_string('standard_offer_template.html', offer_data)

            # Generowanie PDF
            options = {
                'enable-local-file-access': '',
                'quiet': ''
            }
            
            config = pdfkit.configuration(wkhtmltopdf='/usr/bin/wkhtmltopdf')
            pdf = pdfkit.from_string(html_content, False, configuration=config, options=options)

            # Zapisanie PDF do pliku tymczasowego
            file_name = f'Oferta_dla_{data["customer"]["first_name"]}_{data["customer"]["last_name"]}_z_dnia_{current_date}.pdf'
            temp_file_path = f'temp_offer_{client.unique_id}.pdf'
            with open(temp_file_path, 'wb') as temp_file:
                temp_file.write(pdf)

            # **✅ Zapis pliku w modelu ClientFile i powiązanie go z ofertą**
            with open(temp_file_path, 'rb') as file_data:
                client_file = ClientFile.objects.create(
                    client=client,
                    offer=offer,  # ✅ Powiązanie pliku z ofertą
                    file_type='offer',
                    author=request.user,
                    file=File(file_data, name=file_name),
                    description=f"Oferta wygenerowana {current_date}"
                )

            # **✅ Logowanie aktywności dodania pliku oferty**
            if client_file:  # Sprawdzamy, czy plik oferty został utworzony

                # Tworzenie czytelnego opisu aktywności
                client_first_name = client.first_name
                client_last_name = client.last_name
                offer_date = timezone.localtime(client_file.uploaded_at).strftime("%d.%m.%Y")

                formatted_description = f"Oferta dla {client_first_name} {client_last_name} z dnia {offer_date}"

                # Dodanie aktywności do logów
                ActivityLog.objects.create(
                    user=request.user,
                    client=client,
                    action_type="add_offer_file",
                    description=formatted_description
                )
            # **✅ Logowanie aktywności wygenerowania oferty**
            ActivityLog.objects.create(
                user=request.user,
                client=client,
                action_type="generate_offer",
                description=f"Wygenerowano ofertę o wartości {offer.total_price} zł z datą ważności do {valid_until_date}."
            )

            # Dodawanie produktów do oferty
            for product_data in data["customer"]["subsidy_results"]["subsidy_results"]["subsidy_distribution"]:
                product = Product.objects.filter(name=product_data["product_name"]).first()

                if product:
                    offer_product = OfferProduct.objects.create(
                        offer=offer,
                        product=product,
                        quantity=1,
                        price=product_data["product_price"],
                        vat_rate=product_data["vatRate"]
                    )
                    offer_product.required_documents.set(product.required_documents.all())
            previous_milestone = client.milestone

            if client.status == "lead" and client.milestone == "lead":
                client.status = "offer"
                client.milestone = "offer"
                client.save()  # ✅ Zapisujemy zmiany w bazie!
                                # Pobranie aktualnego kamienia milowego przed zmianą

            # **✅ Logowanie aktywności zmiany kamienia milowego**
            if previous_milestone != client.milestone and previous_milestone == "lead":

                ActivityLog.objects.create(
                    user=request.user,
                    client=client,
                    action_type="update_milestone",
                    description=f"Zmieniono etap klienta z '{previous_milestone}' na '{client.milestone}'."
                )
            
            # Usunięcie pliku tymczasowego
            os.remove(temp_file_path)
            # Zwróć PDF jako odpowiedź
            response = HttpResponse(pdf, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{file_name}"'
            return response

        except KeyError as e:
                return JsonResponse({'error': f'Brak klucza w danych: {e}'}, status=400)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Nieprawidłowy format danych JSON.'}, status=400)
        except Exception as e:
            error_details = traceback.format_exc()
            print(f"❌ Unexpected error:\n{error_details}")
            return JsonResponse({'error': f'Nieoczekiwany błąd: {str(e)}', 'details': error_details}, status=500)




@login_required
def calculate_subsidy(request):
    if request.method == "GET":
        try:
            # Odczytaj parametry z URL
            
            program_id = request.GET.get("program")
            option_id = request.GET.get("option")
            products_json = request.GET.get("products")  # Produkty w JSON
            input_value = request.GET.get("input_value")
            max_subsidy = request.GET.get("button_state")
            prefinancing = request.GET.get("prefinancing")
            # Debugowanie wejściowych danych
            heeterExist = False
            isolationExist  = False
            pvExist  = False
            pcExist = False
            mgExist = False
            mcExist = False
            if not products_json:
                return JsonResponse({"error": "Missing products parameter"}, status=400)
            # Sprawdzenie wartości program_id


            # Parsowanie JSON z products
            products = json.loads(products_json)
            subsidy_results = []
            subsidy_distribution = []
            mustPayment = 0
            subsidy_max_amount = 0
            total_gross_price = 0
            total_price = 0
            if not program_id or Decimal(program_id) == 3:


                for product_data in products:
                    product_name = product_data.get("name")
                    grossPrice = (Decimal(product_data.get("price")) * Decimal(product_data.get("vatRate"))  / 100) + Decimal(product_data.get("price"))
                    total_gross_price += grossPrice
                    total_price += Decimal(product_data.get("price"))
                    mustPayment = mustPayment + grossPrice
                    subsidy_distribution.append({
                    "product_name": product_data.get("name"),
                    "attributes": product_data.get("attributes"),
                    "product_price": float(product_data.get("price")),
                    "subsidy": 0,
                    "subsidy_left": 0,
                    "maxSub" : False,
                    "product_id": product_data.get("idProd"),
                    "product_initial_price": product_data.get("initialPrice"),
                    "grossPrice" : float(grossPrice),
                    "vatRate": product_data.get("vatRate")
                    
                })
                subsidy_results =  {
                "subsidy_distribution": subsidy_distribution,
                "total_subsidy": 0,
                "termo_subsidy": 0,
                "heeter_subsidy": 0,
                "mustPayment": float(mustPayment),
                "subsidy_max_amount": 0,
                "total_gross_price": float(total_gross_price),  # 🔹 Konwersja Decimal -> float
                "total_price": float(total_price)  # 🔹 Konwersja Decimal -> float
            }
                result = subsidy_results
                return JsonResponse({"subsidy_results": result})
            # print("Odczytane produkty:", products)
            # Pobierz opcję dotacyjną
            try:
                subsidy_option = SubsidyOption.objects.get(id=option_id)              
            except SubsidyOption.DoesNotExist:
                return JsonResponse({"error": f"Subsidy option with ID {option_id} does not exist."}, status=400)           
            # Iteracja po posortowanych produktach
            for product_data in products:
                product_name = product_data.get("name")
                attributes = product_data.get("attributes", {})
                product_price = Decimal(product_data.get("price"))
                product_id = product_data.get("idProd")
                vatRate = product_data.get("vatRate")
                product_initialPrice = Decimal(product_data.get("initialPrices"))
                # Dopasowanie kryteriów do produktu
                try:
                    product = Product.objects.get(name=product_name)
                    criteria = subsidy_option.product_criteria.filter(product=product).first()
                    if criteria == []:
                        print("brak kryteriów dla: ", product)
                except Product.DoesNotExist:
                    subsidy_results.append({
                        "product_name": product_name,
                        "subsidy_amount": 0,
                        "error": "Product not found in criteria",
                    })
                    continue
                if criteria and str(product.category):
                    # Sprawdź, czy oba kryteria są zdefiniowane
                    if criteria.max_subsidy_percentage and criteria.max_subsidy_amount:
                        try:
                            # Oblicz pierwszą wartość procentową
                            max_percentage_value = round((criteria.max_subsidy_amount / criteria.max_subsidy_percentage) * 100, 2)


                            # Sprawdź, czy produkt przekracza pierwszy próg
                            if criteria.second_max_subsidy_percentage and criteria.second_max_subsidy_amount:

                                if product_price >= max_percentage_value:
                                    temp_price = product_price - max_percentage_value
                                    second_max_value = round((criteria.second_max_subsidy_amount / criteria.second_max_subsidy_percentage) * 100, 2)

                                    if temp_price >= second_max_value:
                                        subsidy_amount = criteria.max_subsidy_amount + criteria.second_max_subsidy_amount
                                    else:
                                        subsidy_amount = criteria.max_subsidy_amount + round((temp_price * criteria.second_max_subsidy_percentage) / 100, 2)
                                else:
                                    subsidy_amount = round((product_price * criteria.max_subsidy_percentage) / 100, 2)
                            elif product_price >= max_percentage_value:
                                subsidy_amount = round(criteria.max_subsidy_amount, 2)
                            else:
                                subsidy_amount = round((product_price * criteria.max_subsidy_percentage) / 100, 2)

                        except Exception as e:
                            print(f"Błąd podczas obliczeń dotacji: {e}")
                            subsidy_amount = 0
                    else:
                        subsidy_amount = 0
                        subsidy_results.append({
                        "name": product_name,
                        "attributes": product_data.get("attributes"),
                        "category": product.category,
                        "price":float(product_price),
                        "subsidy_amount": float(subsidy_amount),
                        "subsidy_percentage": 100,
                        "max_subsidy_amount": 0,
                        'second_subsidy_percentage': 100,
                        "second_max_subsidy_amount": 0,
                        "product_id": product_id,
                        "product_initial_price": product_initialPrice,
                        "vatRate": vatRate
                    })
                    if str(product.category) == "Źródło ogrzewania" and heeterExist:
                        subsidy_amount = 0
                        subsidy_results.append({
                        "name": product_name,
                        "attributes": product_data.get("attributes"),
                        "category": product.category,
                        "price":float(product_price),
                        "subsidy_amount": float(subsidy_amount),
                        "subsidy_percentage": 100,
                        "max_subsidy_amount": 0,
                        'second_subsidy_percentage': 100,
                        "second_max_subsidy_amount": 0,
                        "product_id": product_id,
                        "product_initial_price": product_initialPrice,
                        "vatRate": vatRate
                    })
                else:
                    subsidy_amount = 0
                    subsidy_results.append({
                        "name": product_name,
                        "category": product.category,
                        "attributes": product_data.get("attributes"),
                        "price":float(product_price),
                        "subsidy_amount": float(subsidy_amount),
                        "subsidy_percentage": 100,
                        "max_subsidy_amount": 0,
                        'second_subsidy_percentage': 100,
                        "second_max_subsidy_amount": 0,
                        "product_id": product_id,
                        "product_initial_price": product_initialPrice,
                        "vatRate": vatRate
                    })

                if subsidy_amount != 0:
                    print("Wyniki dofinansowania:", subsidy_amount,"Produkt",str(product_name))
                    if criteria.second_max_subsidy_percentage and criteria.second_max_subsidy_amount:
                        subsidy_results.append({
                            "name": product_name,
                            "category": product.category,
                            "attributes": product_data.get("attributes"),
                            "price":float(product_price),
                            "subsidy_amount": float(subsidy_amount),
                            "subsidy_percentage": float(criteria.max_subsidy_percentage),
                            "max_subsidy_amount": float(criteria.max_subsidy_amount),
                            'second_subsidy_percentage': float(criteria.second_max_subsidy_percentage),
                            "second_max_subsidy_amount": float(criteria.second_max_subsidy_amount),
                            "product_id": product_id,
                            "product_initial_price": product_initialPrice,
                            "vatRate": vatRate
                        })
                    else:
                        subsidy_results.append({
                            "name": product_name,
                            "category": product.category,
                            "attributes": product_data.get("attributes"),
                            "price":float(product_price),
                            "subsidy_amount": float(subsidy_amount),
                            "subsidy_percentage": float(criteria.max_subsidy_percentage),
                            "max_subsidy_amount": float(criteria.max_subsidy_amount),
                            "product_id": product_id,
                            "product_initial_price": product_initialPrice,
                            "vatRate" : vatRate
                        })

                if str(product_name) in ["Instalacja PV - 3 fazy", "Instalacja PV - 1 faza"]:
                    pvExist  = True
                if str(product_name) in ["Magazyn Energii"]:
                    mgExist  = True
                if str(product_name) in ["Magazyn Ciepła"]:
                    mcExist  = True
                if str(product_name) == "Pompa Ciepła":
                    pcExist = True
                if str(product.category) == "Źródło ogrzewania":
                    heeterExist = True
                if str(product.category) == "Izolacja":
                    isolationExist = True
                print(heeterExist,isolationExist,pvExist,pcExist, mgExist)
            if Decimal(program_id) == 2:
                        if pvExist and mgExist:
                            subsidy_max_amount = 23000
                        elif pvExist and mcExist:
                            subsidy_max_amount = 12000
                        elif mgExist:
                            subsidy_max_amount = 16000
                        elif mcExist:
                            subsidy_max_amount = 5000

            if Decimal(program_id) == 1:
                if Decimal(option_id) == 1:
                    print("próg podstawowy")
                    if heeterExist and isolationExist:
                        subsidy_max_amount = subsidy_max_amount + 50000
                        if pvExist:
                            subsidy_max_amount = subsidy_max_amount + 6000
                        if pcExist:
                            subsidy_max_amount = subsidy_max_amount + 10000
                    elif not heeterExist and isolationExist:
                        subsidy_max_amount = subsidy_max_amount + 33000
                    elif pcExist and not isolationExist:
                        subsidy_max_amount = subsidy_max_amount + 27500
                    elif heeterExist and not isolationExist:
                        subsidy_max_amount = subsidy_max_amount + 17200
                        if pvExist:
                            subsidy_max_amount = subsidy_max_amount + 6000
                elif Decimal(option_id) == 2:
                    print("próg podwyższony")
                    if heeterExist and isolationExist:
                        subsidy_max_amount = subsidy_max_amount + 72000
                        if pvExist:
                            subsidy_max_amount = subsidy_max_amount + 9000
                        if pcExist:
                            subsidy_max_amount = subsidy_max_amount + 18000
                    elif not heeterExist and isolationExist:
                        subsidy_max_amount = subsidy_max_amount + 48000
                    elif pcExist and not isolationExist:
                        subsidy_max_amount = subsidy_max_amount + 42400
                    elif heeterExist and not isolationExist:
                        subsidy_max_amount = subsidy_max_amount + 28600
                        if pvExist:
                            subsidy_max_amount = subsidy_max_amount + 9000
                elif Decimal(option_id) == 3:
                    print("próg najwyższy")
                    print("subsidy_max_amount",subsidy_max_amount)  
                    if heeterExist and isolationExist:
                        subsidy_max_amount = subsidy_max_amount + 100000
                        if pvExist:
                            subsidy_max_amount = subsidy_max_amount + 15000
                        if pcExist:
                            subsidy_max_amount = subsidy_max_amount + 20000
                    elif not heeterExist and isolationExist:
                        subsidy_max_amount = subsidy_max_amount + 70000
                    elif pcExist and not isolationExist:
                        subsidy_max_amount = subsidy_max_amount + 55600
                    elif heeterExist and not isolationExist:
                        subsidy_max_amount = subsidy_max_amount + 40800
                        if pvExist:
                            subsidy_max_amount = subsidy_max_amount + 15000
            if max_subsidy == "off":
                result = calculate_optimal_subsidy(subsidy_results, subsidy_max_amount,program_id)
            if max_subsidy == "on":
                result = calculate_maximum_subsidy(subsidy_results, subsidy_max_amount,program_id)
            if prefinancing == "true":
                result["mustPayment"] = result["mustPayment"] - result["total_subsidy"]
                        # Sumowanie cen brutto
            total_gross_price = round(sum(product['grossPrice'] for product in result['subsidy_distribution']), 2)
            total_price = round(sum(product['product_price'] for product in result['subsidy_distribution']), 2)
            result.update({"subsidy_max_amount": subsidy_max_amount,
                           "prefinancing": prefinancing, "total_gross_price": float(total_gross_price), "total_price": float(total_price)})


            print(result, "Cena łączna: ", total_gross_price)
            return JsonResponse({"subsidy_results": result})

        except json.JSONDecodeError:
            return JsonResponse({"error": "Invalid JSON format in 'products'"}, status=400)
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)

    return JsonResponse({"error": "Invalid request method"}, status=405)





@login_required
def get_price(request, product_id):
    attributes_json = request.GET.get("attributes", "{}")
    try:
        attributes = json.loads(attributes_json)  # Dekoduj JSON
    except json.JSONDecodeError:
        return JsonResponse({"success": False, "error": "Nieprawidłowe dane atrybutów."})
    required_fields = {}
    product = get_object_or_404(Product, id=product_id)
    required_fields[product_id] = product.required_fields_offer
    print("Szukany produkt:", product.name)
    print("Przekazane atrybuty:", attributes)
    if product.unit == "m2":
        powierzchnia = extract_number_from_surface(attributes.get("Powierzchnia", None))
        del attributes["Powierzchnia"]
        try:
            configuration = ProductConfiguration.objects.get(product=product, attributes=attributes)
            print("Znaleziono konfigurację:", configuration)
            return JsonResponse({
                "success": True,
                "price": configuration.price*powierzchnia,
                "unit": configuration.unit,
                "required_fields": required_fields
            })
        except ProductConfiguration.DoesNotExist:
            print("Nie znaleziono konfiguracji dla:", attributes)
            return JsonResponse({"success": False, "error": "Cena nie została ustalona."})
    else:
        try:
            configuration = ProductConfiguration.objects.get(product=product, attributes=attributes)
            print("Znaleziono konfigurację:", configuration)
            return JsonResponse({
                "success": True,
                "price": configuration.price,
                "unit": configuration.unit,
                "required_fields": required_fields
            })
        except ProductConfiguration.DoesNotExist:
            print("Nie znaleziono konfiguracji dla:", attributes)
            return JsonResponse({"success": False, "error": "Cena nie została ustalona."})



@login_required
def manage_product_prices(request, product_id):
    product = get_object_or_404(Product, id=product_id)
    configurations = ProductConfiguration.objects.filter(product=product)

    # Pobranie atrybutów z GET lub ustawienie na pusty słownik
    attributes = request.GET.get("attributes", "{}")

    try:
        attributes = json.loads(attributes)  # Parsowanie JSON
    except json.JSONDecodeError:
        attributes = {}  # Jeśli JSON jest niepoprawny, ustaw pusty słownik
    if request.method == "POST":
        if ProductConfiguration.objects.filter(product=product, attributes=attributes).exists():
            return JsonResponse({"error": "Konfiguracja z takimi atrybutami już istnieje."}, status=400)
        form = ProductConfigurationForm(request.POST)
        if form.is_valid():
            configuration = form.save(commit=False)
            configuration.product = product
            if product.unit == "m2":
                del attributes["Powierzchnia"]
            configuration.attributes = attributes
            configuration.save()
            return redirect("manage_product_prices", product_id=product.id)
    else:
        form = ProductConfigurationForm()

    return render(request, "manage_product_prices.html", {
        "product": product,
        "configurations": configurations,
        "form": form,
        "attributes": json.dumps(attributes, indent=2),  # JSON do wyświetlenia w textarea
    })



    
@login_required
def edit_product_configuration(request, config_id):
    configuration = get_object_or_404(ProductConfiguration, id=config_id)
    if request.method == "POST":
        form = ProductConfigurationForm(request.POST, instance=configuration)
        if form.is_valid():
            form.save()
            return redirect("manage_product_prices", product_id=configuration.product.id)
    else:
        form = ProductConfigurationForm(instance=configuration)

    return render(request, "edit_product_configuration.html", {
        "form": form,
        "configuration": configuration,
    })
        
@login_required
def delete_product_configuration(request, config_id):
    configuration = get_object_or_404(ProductConfiguration, id=config_id)
    product_id = configuration.product.id
    configuration.delete()
    return redirect("manage_product_prices", product_id=product_id)


@login_required
def get_selected_products(request):
    selected_products = [
        {"name": "Piec na pelet", "price": 10000, "vat_rate": 23},
        {"name": "Ruszt żeliwny", "price": 500, "vat_rate": 8},
    ]
    return JsonResponse(selected_products, safe=False)



@login_required
def get_products_by_category(request, category_id):
    try:
        products = Product.objects.filter(category_id=category_id)
        products_data = [
            {"id": product.id, "name": product.name, "description": product.description}
            for product in products
        ]
        return JsonResponse({"success": True, "products": products_data})
    except Product.DoesNotExist:
        return JsonResponse({"success": False, "error": "Produkty nie znalezione."})

@login_required
def get_product_details(request, product_id):
    product = get_object_or_404(Product, id=product_id)
    attributes = ProductAttribute.objects.filter(product=product)

    data = {
        "success": True,
        "product": {"name": product.name},
        "attributes": [
            {
                "name": attr.name,
                "type": attr.input_type,
                "placeholder": attr.placeholder if attr.input_type == "text" else None,
                "options": attr.options.split(",") if attr.input_type == "dropdown" else [],
            }
            for attr in attributes
        ],
    }
    return JsonResponse(data)






@login_required
def lead_form(request):
    leads = Client.objects.all().order_by('-potential')  # Sortowanie według potencjału
    form = ClientForm()

    if request.method == 'POST':
        form = ClientForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('klienci')  # Przekierowanie po zapisaniu formularza

    return render(request, 'lead_form.html', {'form': form, 'leads': leads})

@login_required
def get_lead_data(request, lead_id):
    lead = get_object_or_404(Client, id=lead_id)
    data = {
            "first_name": lead.first_name,
            "last_name": lead.last_name,
            "phone": lead.phone,
            "email": lead.email,
            "street": lead.street,
            "house_number": lead.house_number,
            "city": lead.city,
            "postal_code": lead.postal_code,
            "postal": lead.postal,
            "current_heating_source": lead.current_heating_source,
            "construction_permission_year": lead.construction_permission_year,
            "land_registry_number": lead.land_registry_number,
            "plot_number": lead.plot_number,
            "income_threshold": lead.income_threshold,
            "household_members": lead.household_members,
            "farm_conversion_hectares": lead.farm_conversion_hectares,
            "income_per_person": lead.income_per_person,
            "marital_status": lead.marital_status,
            "is_sole_owner": lead.is_sole_owner,
            "runs_business": lead.runs_business,
            "has_joint_property": lead.has_joint_property,
            "potential": lead.potential,
            "person_type": lead.person_type,
            "has_different_investment_address": lead.has_different_investment_address,
            "investment_street": lead.investment_street,
            "investment_house_number": lead.investment_house_number,
            "investment_postal_code": lead.investment_postal_code,
            "status_record": lead.status,
            "unique_id": lead.unique_id,
    }
    return JsonResponse({"success": True, "lead": data})
    
@login_required
def add_client(request):
    form = ClientForm(request.POST or None)

    if request.method == "POST":
        print("Dane formularza:", request.POST)  # Debugowanie danych
        print("Błędy formularza:", form.errors)  # Debugowanie błędów walidacji

        if form.is_valid():
            client = form.save(commit=False)
            client.user = request.user
            client.save()
            return redirect("leady")  # Przekierowanie po zapisaniu
        else:
            messages.error(request, "Formularz zawiera błędy. Popraw je i spróbuj ponownie.")

    return render(request, "add_client.html", {"form": form})

@login_required
def edit_client(request, client_id):
    client = get_object_or_404(Client, id=client_id, user=request.user)  # Dodajemy filtrowanie po użytkowniku
    if request.method == 'POST':
        if 'delete' in request.POST:  # Sprawdzamy, czy kliknięto przycisk "Usuń"
            client.delete()
            return redirect('leady')  # Po usunięciu przekierowanie na listę klientów
                # Konwersja request.POST na edytowalną kopię, bo jest immutable
        post_data = request.POST.copy()

        # Jeśli milestone nie został wysłany w formularzu, przypisz istniejącą wartość
        if 'milestone' not in post_data:
            post_data['milestone'] = client.milestone

        # Tworzenie formularza na podstawie zmodyfikowanych danych
        form = ClientForm(post_data, instance=client)
        if form.is_valid():
            client.modified_by = request.user
            client.save()
            previous_data = {field: getattr(client, field) for field in form.cleaned_data}
            form.save()

            changes = []
            for field, new_value in form.cleaned_data.items():
                old_value = previous_data[field]
                if old_value != new_value:
                    changes.append(f"{field.replace('_', ' ')}: {old_value} ➝ {new_value}")

            if changes:
                ActivityLog.objects.create(
                    user=request.user,
                    client=client,
                    action_type="edit_client",
                    description=f"Zaktualizowano dane klienta: {', '.join(changes)}"
                )
            return redirect('oferty')  # Po edycji przekierowanie na listę klientów
    else:
        form = ClientForm(instance=client)
    
    return render(request, 'edit_client.html', {'form': form, 'client': client})


@login_required
def client_card(request, client_id):
        # Sprawdzamy, czy użytkownik jest w grupie "biuro"
    if request.user.groups.filter(name="biuro").exists():
        # Jeśli użytkownik jest w biurze, ma dostęp do wszystkich klientów
        client = get_object_or_404(Client, id=client_id)
    else:
        # Handlowiec widzi tylko klientów, których dodał
        client = get_object_or_404(Client, id=client_id, user=request.user)
    final_offer = client.final_offer  # Pobranie finalnej oferty klienta, jeśli istnieje
    active_tab = request.GET.get("tab", request.POST.get("tab", "overview"))  # Domyślnie "Przegląd"
    file_type_label = ClientFile.FILE_TYPE_CHOICES[0][1]  # Pobranie pierwszej wartości
        # Pobranie niewykonanych zadań (maksymalnie 3)
    pending_tasks = client.tasks.filter(completed=False).order_by("-due_date")[:3]

    # Jeśli mamy mniej niż 3 niewykonane zadania, dobieramy wykonane
    if pending_tasks.count() < 3:
        completed_tasks = client.tasks.filter(completed=True).order_by("-due_date")[:3 - pending_tasks.count()]
    else:
        completed_tasks = []

    # Łączymy listy
    recent_tasks = list(pending_tasks) + list(completed_tasks)
    required_documents_list = []
    # Pobranie rzeczywistych ofert (z modelu Offer)
    offers = Offer.objects.filter(client=client).prefetch_related("products", "products__required_documents")
        # Lista kroków projektu
    all_milestones = [
        {"key": "lead", "label": "Lead"},
        {"key": "offer", "label": "Oferta"},
        {"key": "agreement", "label": "Umowa"},
        {"key": "prefinancing_request", "label": "I transza"},
        {"key": "realization", "label": "Realizacja"},
        {"key": "second_tranche_request", "label": "II transza"},
        {"key": "completed", "label": "Zakończono"},
    ]

    # Filtracja kroków w zależności od prefinansowania
    if not client.prefinancing:
        milestones = [step for step in all_milestones if step["key"] in ["lead", "offer", "agreement", "realization", "completed"]]
    else:
        milestones = all_milestones

# Iteracja przez oferty i produkty
    # for offer in offers:
    #     print(f"Oferta: {offer.id} dla klienta: {offer.client}, cena całkowita: {offer.total_price}, całkowita kowta do zapłaty: {offer.must_payment}")
    #     print(client.phone)
    #     for product in offer.products.all():
    #         print(f"  - Produkt: {product.product.name}, Ilość: {product.quantity}, Cena: {product.price}, Cena brutto: {Decimal(product.price)* (1 + Decimal(product.vat_rate) / 100)}, Vat: {product.vat_rate}%")


    # Pobranie plików ofertowych powiązanych z klientem
    offer_files = ClientFile.objects.filter(client=client, file_type="offer")

    # Pobranie wymaganych dokumentów dla każdej oferty (powiązanych z produktami)
    required_documents = {
        offer.id: list(set(
            doc for product in offer.products.all() for doc in product.required_documents.all()
        ))
        for offer in offers
    }

    for item in list(required_documents.values()):
        temp = []
        for i in item:
           temp.append(i.name)
        required_documents_list.append(temp)
        
    for offer in offers:
        document_set = set()  # Używamy set, aby uniknąć duplikatów

        for product in offer.products.all():
            for doc in product.required_documents.all():
                document_set.add(doc)  # Dodajemy pełny obiekt `DocumentTemplate`
        
        required_documents[offer.id] = list(document_set)  # Zamieniamy na listę przed przypisaniem
    #Połączenie oferty z wymaganymi dokumentami
    offers_with_documents = zip(offers, required_documents_list)

    # print("required_documents",list(required_documents.values()))
    # Pobieranie spotkań
    meetings = client.meetings.order_by("occurred", "note", "meeting_date")
    meetings_rest = client.meetings.order_by("meeting_date")

    # Obsługa dodawania spotkania
    if request.method == "POST" and "add_meeting" in request.POST:
        meeting_form = MeetingForm(request.POST)
        if meeting_form.is_valid():
            meeting = meeting_form.save(commit=False)
            meeting.client = client
            meeting.author = request.user
            meeting.save()
            messages.success(request, "Spotkanie zostało dodane.")
            # Logowanie aktywności
            # Pobranie godziny i daty spotkania w odpowiednim formacie
            meeting_time = timezone.localtime(meeting.meeting_date).strftime("%H:%M")
            meeting_date = timezone.localtime(meeting.meeting_date).strftime("%d.%m.%Y")

            # Tworzenie opisu aktywności
            description = f" o {meeting_time} dnia {meeting_date} z {client.first_name} {client.last_name}"

            # Logowanie aktywności
            ActivityLog.objects.create(
                user=request.user,
                client=client,
                action_type="add_meeting",
                description=description
            )
            return redirect(f"{request.path}?tab=appointments")  # Przekierowanie po dodaniu
    else:
        meeting_form = MeetingForm()

    # Podział plików na kategorie i sortowanie malejąco po dacie utworzenia
    files_by_category = {
        label: client.files.filter(file_type=category).order_by('-uploaded_at')
        for category, label in ClientFile.FILE_TYPE_CHOICES
    }
        
    # Pobranie 10 ostatnio dodanych plików
    recent_files = client.files.order_by("-uploaded_at")[:10]
    
    # Grupowanie plików po kategoriach
    file_category_counts = defaultdict(int)
    for file in recent_files:
        file_category_counts[file.get_file_type_display()] += 1  # Pobranie czytelnej nazwy kategorii

    if request.method == "POST" and "add_file" in request.POST:
            files = request.FILES.getlist('files')
            file_type = request.POST.get("file_type")
            description = request.POST.get("description")
            file_type_label = dict(ClientFile.FILE_TYPE_CHOICES).get(file_type, "Nieznany typ pliku")   
            if not files:
                return JsonResponse({"success": False, "error": "Nie wybrano plików!"})

            saved_files = []
            for file in files:
                client_file = ClientFile(
                    client=client,
                    author=request.user,
                    file=file,
                    file_type=file_type,
                    description=description,
                )
                client_file.save()
                saved_files.append(client_file)
            
            if saved_files:
                # Logowanie aktywności jako jedno zdarzenie
                file_count = len(saved_files)
                file_word = "plik" if file_count == 1 else "pliki"
                ActivityLog.objects.create(
                    user=request.user,
                    client=client,
                    action_type="add_file",
                    description=f"Dodał/a {file_count} {file_word} ({file_type_label})"
                )

                return JsonResponse({"success": True, "file_type": file_type_label})

            return JsonResponse({"success": False, "error": "Błąd podczas zapisu plików."})




    # Pobieranie zadań - niewykonane jako pierwsze
    tasks = client.tasks.order_by("completed", "due_date")

    # Obsługa dodawania nowego zadania
    if request.method == "POST" and "add_task" in request.POST:
        task_form = TaskForm(request.POST)
        if task_form.is_valid():
            new_task = task_form.save(commit=False)
            new_task.client = client
            new_task.author = request.user
            new_task.save()
            messages.success(request, "Zadanie zostało dodane.")
            # Logowanie aktywności
            ActivityLog.objects.create(
                user=request.user,
                client=client,
                action_type="add_task",
                description=f"{new_task.text[:50]}..." if len(new_task.text) > 50 else f"{new_task.text}"
            )
            # Powiadamiamy tylko odpowiedniego użytkownika
            create_notification("zadanie", request.user, client_id=client.id)
            return redirect(f"{request.path}?tab=tasks")  # Przekierowanie po dodaniu
    else:
        task_form = TaskForm()

    # Pobieranie notatek - ważne jako pierwsze
    notes = client.notes.order_by("-is_important", "-added_date")

    # Obsługa dodawania notatek
    if request.method == "POST" and "add_note" in request.POST:
        note_form = NoteForm(request.POST)
        if note_form.is_valid():
            note = note_form.save(commit=False)
            note.client = client
            note.author = request.user
            note.save()
            messages.success(request, "Notatka została dodana!")
                    # Logowanie aktywności
            ActivityLog.objects.create(
            user=request.user,
            client=client,
            action_type="add_note",
            description=f"{note.text[:50]}{'...' if len(note.text) > 50 else ''}"  # Skrócony opis notatki
        )
            # Powiadamiamy tylko odpowiedniego użytkownika
            create_notification("notatkę", request.user, client_id=client.id)
            return redirect(f"{request.path}?tab=notes")  # Przekierowanie po dodaniu
    else:
        note_form = NoteForm()
        
    # Pobieranie płatności - ważne jako pierwsze
    payments = client.payments.select_related("invoice_file").order_by("-due_date")  # Optymalizacja zapytań do bazy

    if request.method == "POST" and "add_payment" in request.POST:
        payment_form = PaymentForm(request.POST, request.FILES)

        if payment_form.is_valid():
            payment = payment_form.save(commit=False)
            payment.client = client

            if payment.status == "paid":
                payment.payment_date = date.today()  # ✅ Automatyczna data, jeśli status to "Opłacona"

            # Obsługa dodania faktury
            invoice_file = request.FILES.get("invoice")
            if invoice_file:
                invoice_entry = ClientFile.objects.create(
                    client=client,
                    file_type="invoice",
                    file=invoice_file,
                    author=request.user,
                    description=f"Faktura za płatność {payment.amount} zł z dnia {payment.due_date.strftime('%d.%m.%Y')}",
                )
                # Powiadamiamy tylko odpowiedniego użytkownika
                create_notification("płatność", request.user, client_id=client.id)
                payment.invoice_file = invoice_entry  # ✅ Przypisanie faktury jako obiekt, nie URL!

            payment.save()  # ✅ Ważne! Musisz zapisać ponownie po przypisaniu faktury
             # Logowanie aktywności płatności
            ActivityLog.objects.create(
            user=request.user,
            client=client,
            action_type="add_payment",
            description=f"{payment.amount} zł, termin {payment.due_date.strftime('%d.%m.%Y')}"
        )

            # Logowanie aktywności faktury (jeśli została dodana)
            if invoice_file:
                ActivityLog.objects.create(
                    user=request.user,
                    client=client,
                    action_type="add_file",
                    description=f"faktura - kwota: {payment.amount} zł z dnia {payment.due_date.strftime('%d.%m.%Y')}"
                )
            messages.success(request, "Płatność została dodana.")
            return redirect(request.path + "?tab=payment")  # ✅ Pozostanie na zakładce "Płatności"

    # Formularz do dodawania płatności
    else:
        payment_form = PaymentForm()
        
    # 🆕 Dodawanie faktury do istniejącej płatności
    if request.method == "POST" and "add_invoice" in request.POST:
        payment_id = request.POST.get("payment_id")
        payment = get_object_or_404(Payment, id=payment_id, client=client)

        invoice_file = request.FILES.get("invoice")
        if invoice_file:
            invoice_entry = ClientFile.objects.create(
                client=client,
                file_type="invoice",
                file=invoice_file,
                author=request.user,
                description=f"Faktura za płatność {payment.amount} zł z dnia {payment.due_date.strftime('%d.%m.%Y')}",
            )
            ActivityLog.objects.create(
                user=request.user,
                client=client,
                action_type="add_file",
                description=f"faktura - kwota: {payment.amount} zł z dnia {payment.due_date.strftime('%d.%m.%Y')}"
            )
            payment.invoice_file = invoice_entry
            payment.save()
            messages.success(request, "Faktura została dodana do płatności.")

        return redirect(request.path + "?tab=payment")

    if request.method == "POST" and "mark_as_paid" in request.POST:
        payment_id = request.POST.get("payment_id")
        payment = get_object_or_404(Payment, id=payment_id, client=client)
        
        if payment.status != "paid":  # ✅ Sprawdzamy, czy już nie jest oznaczona jako opłacona
            payment.status = "paid"
            payment.payment_date = date.today()
            payment.save()

            # Logowanie aktywności oznaczenia jako opłacona
            ActivityLog.objects.create(
                user=request.user,
                client=client,
                action_type="mark_as_paid",
                description=f"jako opłacona: {payment.amount} zł z dnia {payment.due_date.strftime('%d.%m.%Y')} jako opłaconą."
            )
            # Powiadamiamy tylko odpowiedniego użytkownika
            create_notification("dodało status 'opłacona' do płatności", request.user, client_id=client.id)
            messages.success(request, "Płatność została oznaczona jako opłacona.")
        else:
            messages.warning(request, "Płatność już wcześniej była oznaczona jako opłacona.")

        return redirect(request.path + "?tab=payment")
    else:
        payment_form = PaymentForm()


    if request.method == "POST"and "generateDoc" in request.POST:
            selected_templates =  DocumentTemplate.objects.all()
            selected_offer_id = request.POST.get("selected_offer", "").strip()
            has_representative = request.POST.get("has_representative")
            representative_name = request.POST.get("representative_name")
            client_type = request.POST.get("client_type")
            
            if not has_representative:
                representative_name = ""
            
            
                

            # Pobranie wybranej oferty
            selected_offer = get_object_or_404(Offer, id=selected_offer_id)
              # Pobranie produktów powiązanych z ofertą
            products = OfferProduct.objects.filter(offer=selected_offer).select_related("product")
            
            required_documents = set()
            for product in selected_offer.products.all():
                required_documents.update(product.required_documents.all())
            #Tworze listę z wszystkimi placeholderami w wymaganych dokumentach
            # Tworzymy jedną listę z wszystkimi placeholderami w wymaganych dokumentach
            all_placeholders_in_required_documents = [
                placeholder for item in required_documents for placeholder in item.placeholders
            ]
            price_product_details = {}
            # Pobranie wymaganych placeholderów tylko z ProductDocumentRequirement
            all_required_placeholders = set()
            for product in products:
                product_requirements = ProductDocumentRequirement.objects.filter(product=product.product)
                print(product.product.name)
                gross_price = ((product.vat_rate * product.price) / 100) + product.price
                if "pv" in product.product.name.lower():  # Konwersja na małe litery dla bezpieczeństwa
                    price_product_details.update({"wynagrodzenie netto za instalację fotowoltaiczną": product.price, "wynagrodzenie brutto za instalację fotowoltaiczną": gross_price, "słownie wynagrodzenie za instalację fotowoltaiczną": kwota_slownie(gross_price)})
                if "magazyn" in product.product.name.lower():  # Konwersja na małe litery dla bezpieczeństwa
                    price_product_details.update({"wynagrodzenie netto za magazyn energii": product.price, "wynagrodzenie brutto za magazyn energii": gross_price, "słownie wynagrodzenie za magazyn energii": kwota_slownie(gross_price)})
                     
                for requirement in product_requirements:
                    all_required_placeholders.update(requirement.required_placeholders)  # Dodajemy tylko wymagane placeholdery



            # Pobranie wybranej metody płatności
            payment_method = request.POST.get("payment-method")

            # Pełny słownik wartości (wszystkie opcje domyślnie puste)
            payment_data_to_template = {
                "przelew a I transza 50%": "",
                "przelew a II transza 50%": "",
                "przelew b I transza 30%": "",
                "przelew b II transza": "",
                "przelew b III transza 20%": "",
                "przelew d I transza": "",
                "przelew d II transza": "",
                "przelew d III transza": "",
                "przelew d kredyt": "",
                "przelew d": "",
                "I transza całkowitej inwestycji": "",
                "data wpłaty całkowitej inwestycji": "",
            }

            # Obsługa różnych metod płatności
            if payment_method == "transfer-a":
                payment_data_to_template["przelew a I transza 50%"] = request.POST.get("transfer-a1", "")
                payment_data_to_template["przelew a II transza 50%"] = request.POST.get("transfer-a2", "")

            elif payment_method == "transfer-b":
                payment_data_to_template["przelew b I transza 30%"] = request.POST.get("transfer-b1", "")
                payment_data_to_template["przelew b II transza"] = request.POST.get("transfer-b2", "")
                payment_data_to_template["przelew b III transza 20%"] = request.POST.get("transfer-b3", "")

            elif payment_method == "split-payment":
                payment_data_to_template["przelew d I transza"] = request.POST.get("split-t1", "")
                payment_data_to_template["przelew d II transza"] = request.POST.get("split-t2", "")
                payment_data_to_template["przelew d III transza"] = request.POST.get("split-t3", "")
                payment_data_to_template["przelew d kredyt"] = request.POST.get("split-credit", "")
                payment_data_to_template["przelew d"] = request.POST.get("split-transfer", "")

            elif payment_method == "clean-air":
                payment_data_to_template["I transza całkowitej inwestycji"] = selected_offer.must_payment
                payment_data_to_template["data wpłaty całkowitej inwestycji"] = request.POST.get("clean-air-date", "")

            payment_details = {}

            payment_details = {
                "transfer_a1": request.POST.get("transfer-a1", ""),
                "transfer_a2": request.POST.get("transfer-a2", ""),
                "transfer_b1": request.POST.get("transfer-b1", ""),
                "transfer_b2": request.POST.get("transfer-b2", ""),
                "transfer_b3": request.POST.get("transfer-b3", ""),
                "split_credit": request.POST.get("split-credit", ""),
                "split_transfer": request.POST.get("split-transfer", ""),
                "split_t1": request.POST.get("split-t1", ""),
                "split_t2": request.POST.get("split-t2", ""),
                "split_t3": request.POST.get("split-t3", ""),
                "clean_air_date": request.POST.get("clean-air-date", ""),
            }





            saved_files = []
            for template in selected_templates:
                doc_path = template.file.path
                doc = docx.Document(doc_path)
                        # Dane do podstawienia w szablonie
                # Pobranie profilu użytkownika
            profile, created = Profile.objects.get_or_create(
                user=request.user,
                defaults={"current_contract_number": "1"}  # ✅ Przypisanie domyślnej wartości
            )

            
            # Pobranie aktualnej daty
            today = datetime.today()
            day = today.strftime("%d")  # dzień (np. "01")
            month = today.strftime("%m")  # miesiąc (np. "02")
            year = today.strftime("%y")  # ostatnie dwie cyfry roku (np. "25")


  
            # Tworzenie numeru umowy
            contract_number = f"{profile.current_contract_number}/{day}/{month}/{year}/{request.user.first_name[0]}{request.user.last_name[0]}"

            # Tworzymy słownik dostępnych danych
            available_data = {
                # Informacje o umowie
                "numer umowy": contract_number,
                "data umowy": selected_offer.created_at.strftime("%d.%m.%Y"),

                # Dane osobowe klienta
                "imię i nazwisko klienta": f"{client.first_name} {client.last_name}",
                "imię i nazwisko handlowca": f"{request.user.first_name} {request.user.last_name}",
                "imię i nazwisko reprezentanta": f"{client.user.first_name} {client.user.last_name}" if client.user else "",
                "Data urodzenia": client.birth_date,
                "pesel": client.pesel if client.pesel and client.pesel.strip() != "" else None,
                "nip": client.nip,
                "numer dowodu osobistego": client.id_card_number if client.id_card_number and client.id_card_number.strip() != "" else None,
                "numer telefonu": client.phone if client.phone  and client.phone.strip() != "" else None,
                "e-mail": client.email if client.email and client.email.strip() != "" else None,
                "imię i nazwisko reprezentanta": representative_name,

                # Adres klienta
                "ulica": client.street if client.street and client.street.strip() != "" else None,
                "numer domu": client.house_number if client.house_number and client.house_number.strip() != "" else None,
                "miejscowość": client.city if client.city and client.city.strip() != "" else None,
                "kod pocztowy": client.postal_code if client.postal_code and client.postal_code.strip() != "" else None,
                "poczta": client.postal if client.postal and client.postal.strip() != "" else None,
                "województwo": przypisz_wojewodztwo(str(client.postal_code)),  # Brakuje pola w modelu, można dodać, jeśli istnieje w innym miejscu

                # Adres inwestycji (jeśli inny)
                "ulica (adres inwestycji)": client.investment_street if client.has_different_investment_address else client.street,
                "numer domu (adres inwestycji)": client.investment_house_number if client.has_different_investment_address else client.house_number,
                "kod pocztowy (adres inwestycji)": client.investment_postal_code if client.has_different_investment_address else client.postal_code,

                # Informacje finansowe i program Czyste Powietrze
                "całkowita kwota inwestycji": selected_offer.total_price,
                "kwota prefinansowania przysługująca dla Inwestor": "Tak" if client.prefinancing else "Nie",
                "I transza całkowitej inwestycji": selected_offer.must_payment if hasattr(selected_offer, "must_payment") else "",
                "data wpłaty całkowitej inwestycji": "",  # Może być przekazana dynamicznie

                # Status klienta
                "status klienta": dict(Client.STATUS_CHOICES).get(client.status, ""),
                "potencjał klienta": dict(Client.POTENTIAL_CHOICES).get(client.potential, ""),
                "etap klienta": dict(Client.MILESTONE_CHOICES).get(client.milestone, ""),

                # Informacje prawne i działalność
                "stan cywilny": dict(Client.marital_status.field.choices).get(client.marital_status, ""),
                "czy jest jedynym właścicielem": "Tak" if client.is_sole_owner else "Nie",
                "czy prowadzi działalność gospodarczą": "Tak" if client.runs_business else "Nie",
                "czy posiada wspólność majątkową": "Tak" if client.has_joint_property else "Nie",
                "czy działalność prowadzona pod adresem inwestycji": "Tak" if client.business_at_investment_address else "Nie",

                # Informacje o budynku
                "rok pozwolenia na budowę": client.construction_permission_year if client.construction_permission_year else "",
                "numer księgi wieczystej": client.land_registry_number if client.land_registry_number else "",
                "numer działki": client.plot_number if client.plot_number else "",

                # Próg dochodowy i gospodarstwo domowe
                "próg dochodowy": dict(Client.INCOME_CHOICES).get(client.income_threshold, ""),
                "liczba osób w gospodarstwie domowym": client.household_members if client.household_members else "",
                "ilość hektarów przeliczeniowych": client.farm_conversion_hectares if client.farm_conversion_hectares else "",
                "dochód na osobę": client.income_per_person if client.income_per_person else "",

                # Informacje o ogrzewaniu
                "aktualne źródło ogrzewania": client.current_heating_source if client.current_heating_source else "",
            }
            print((available_data["miejscowość"]), type(available_data["miejscowość"]))
            if selected_offer.additional_terms:
                available_data["brak ustaleń imię i nazwisko klienta"] = ""
                available_data["ustalenia imię i nazwisko klienta"] = f"{client.first_name} {client.last_name}"
                available_data["dodatkowe ustalenia"] = selected_offer.additional_terms
            else:
                available_data["brak ustaleń imię i nazwisko klienta"] = f"{client.first_name} {client.last_name}"
                available_data["ustalenia imię i nazwisko klienta"] = " "
                available_data["dodatkowe ustalenia"] = " "
            # Jeśli klient ma użytkownika przypisanego, pobierz dodatkowe dane
            if client_type == "individual":
                available_data["nip"] =""
                available_data["krs"] =""
            else:
                available_data["pesel"] =""
            print("asas", available_data)
            # Dodanie danych płatności do available_data
            available_data.update(payment_data_to_template)
            available_data.update(price_product_details)
            # Tworzymy kontekst tylko dla wymaganych placeholderów
            context = {key: available_data[key] for key in all_required_placeholders if key in available_data}

        
            filled_data = {}
            # **✅ 1️⃣ Sprawdzamy brakujące dane**
            missing_fields = [field for field in all_required_placeholders if context.get(field) is None or ""]

            # Tworzymy kopię listy, aby móc bezpiecznie modyfikować missing_fields
            missing_fields_copy = missing_fields[:]

            filled_data = {}  # Słownik do przechowywania uzupełnionych wartości

            for item in missing_fields_copy:
                value = request.POST.get(item, "").strip()  # Pobranie wartości i usunięcie zbędnych spacji
                
                if value:  # Jeśli wartość nie jest pusta
                    filled_data[item] = value  # Dodajemy do słownika
                    missing_fields.remove(item)  # Usuwamy uzupełnione pole z oryginalnej listy

            # **✅ 2️⃣ Uzupełniamy dane w bazie danych**
            updated = False  # Flaga zmiany w bazie
            FIELD_MAPPING = {
                "ulica": "street",
                "miejscowość": "city",
                "kod pocztowy": "postal_code",
                "numer dowodu osobistego": "id_card_number",
                "pesel": "pesel",
                "nip": "nip",
                "poczta": "postal",
                "e-mail": "email",
            }

            for field, value in filled_data.items():
                if field in FIELD_MAPPING:  # Sprawdzamy, czy pole ma swoje odwzorowanie w modelu
                    model_field = FIELD_MAPPING[field]  # Pobieramy odpowiadające pole w modelu
                    if hasattr(client, model_field):  # Sprawdzamy, czy klient ma to pole
                        setattr(client, model_field, value)  # Ustawiamy nową wartość
                        updated = True  # Oznaczamy, że klient został zaktualizowany

            if updated:
                client.save()  # **Zapisujemy zmiany w bazie**
                print(f"✅ Uzupełniono brakujące pola w bazie dla klienta {client.first_name} {client.last_name}: {filled_data}")

            # **✅ 3️⃣ Aktualizacja kontekstu dla wyświetlenia brakujących pól**
            available_data.update(filled_data)
            if missing_fields:
                context = {
        "client": client,
        "final_offer": final_offer,
        "active_tab": active_tab,
        "files_by_category": files_by_category,
        "recent_file_categories": dict(file_category_counts),
        "recent_notes": client.notes.order_by("-added_date")[:3],
        "recent_tasks": recent_tasks,
        "recent_activities": client.activities.order_by("-timestamp")[:6],
        "milestones": milestones,
        "task_form": task_form,
        "meeting_form": meeting_form,
        "note_form": note_form,
        "tasks": tasks,
        "meetings": meetings,
        "meetings_rest": meetings_rest,
        "notes": notes,
        "active_type": file_type_label,
        "payments": payments,
        "payment_form": payment_form,
        "offers_with_documents": offers_with_documents,
        "offers": offers,
        "missing_fields": missing_fields,
        "payment_method":payment_method,
        "payment_details": payment_details,
        "offer_ids": selected_offer_id.strip()
    }

                return render(request, "client_card.html", context)
            empty_placeholders = {}

            # Tworzymy kopię, aby uniknąć modyfikacji oryginalnej listy podczas iteracji
            all_placeholders_copy = set(all_placeholders_in_required_documents)
            
            # Iterujemy po wszystkich placeholderach w dokumentach
            for placeholder in all_placeholders_copy:
                # Jeśli placeholder nie jest wymagany, ustawiamy pusty string
                if placeholder not in all_required_placeholders:
                    empty_placeholders[placeholder] = ""

            # Aktualizacja dostępnych danych
            available_data.update(empty_placeholders)
            # Tworzymy kontekst tylko dla wymaganych placeholderów
            context = {key: available_data.get(key, "") for key in all_placeholders_in_required_documents}
            # Tworzymy pełną kopię available_data, aby nie gubić danych
            context = available_data.copy()
        
            # Dodajemy puste wartości dla placeholderów, które nie są w available_data
            for placeholder in all_placeholders_in_required_documents:
                if placeholder not in context:
                    context[placeholder] = ""
            context = {key.strip(): value for key, value in context.items()}
            print("Kontekst dla dokumentów:", context)
            print("Czy klucz istnieje:", "słownie wynagrodzenie za okna" in context)
            saved_files = []  # Lista zapisanych plików

            for template in required_documents:  # Iteracja po wielu szablonach
                doc = Document(template.file.path)  # Otwórz dokument na podstawie szablonu

                # Zamiana placeholderów w zwykłych paragrafach
                for para in doc.paragraphs:
                    for key, value in context.items():
                        placeholder = f"{{{{ {key} }}}}"
                        if placeholder in para.text:
                            para.text = para.text.replace(placeholder, str(value))

                # Zamiana placeholderów w tabelach
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for key, value in context.items():
                                placeholder = f"{{{{ {key} }}}}"
                                if placeholder in cell.text:
                                    cell.text = cell.text.replace(placeholder, str(value))

                # Tworzenie folderu dla plików klienta
                client_dir = os.path.join(settings.MEDIA_ROOT, "client_files", str(selected_offer.client.id))
                os.makedirs(client_dir, exist_ok=True)

                # Tworzenie unikalnej nazwy dla pliku
                filename = f"{template.name}_{client.first_name.replace(' ', '_')}.docx"
                file_path = os.path.join(client_dir, filename)

                # Zapisujemy plik
                doc.save(file_path)

                # Tworzymy opis dla dokumentu
                # Wyciągnięcie nazwy dokumentu bez imienia i rozszerzenia
                document_name = filename.rsplit("_", 1)[0]  # Usuwa wszystko po ostatnim "_"

                # Tworzenie opisu dla dokumentu
                description = f"{document_name} dla {client.first_name} {client.last_name}, numer umowy {contract_number}, utworzony {today.strftime('%d.%m.%Y')}"


                # Zapisujemy plik w bazie danych
                client_file = ClientFile.objects.create(
                    client=selected_offer.client,
                    author=request.user,
                    description=description,
                    file=f"client_files/{selected_offer.client.id}/{filename}",
                    file_type=template.document_type  # Przypisanie typu dokumentu
                )
                saved_files.append(client_file)  # ✅ Zamiast URL zapisujemy cały obiekt ClientFile

            # **✅ Logowanie aktywności generowania dokumentów**
            if saved_files:
                ActivityLog.objects.create(
                    user=request.user,
                    client=client,
                    action_type="generate_document",
                    description=f"Wygenerowano dokumenty: {', '.join([doc.name for doc in selected_templates])}, numer umowy {contract_number}."
                )

            # **✅ Logowanie aktywności dodania plików**
            if saved_files:
                # Mapowanie kodów typów plików na ich polskie nazwy
                file_type_dict = dict(ClientFile.FILE_TYPE_CHOICES)

                # Pobranie unikalnych kategorii plików i ich polskich nazw
                file_types = {file_type_dict.get(file.file_type, "Nieznany typ") for file in saved_files}

                ActivityLog.objects.create(
                    user=request.user,
                    client=client,
                    action_type="add_file",
                    description=f"Dodano pliki: {', '.join(file_types)} ({len(saved_files)})."
                )
            # Po zakończeniu przetwarzania wszystkich plików, przekierowanie
            return redirect(request.path + "?tab=overview")




    # Oznaczenie statusu kroków
    for step in milestones:
        current_index = milestones.index(step)
        milestone_index = next((i for i, s in enumerate(milestones) if s["key"] == client.milestone), -1)
        if current_index < milestone_index:
            step["status"] = "completed"
        elif current_index == milestone_index:
            step["status"] = "active"
        else:
            step["status"] = "upcoming"

    context = {
        "client": client,
        "final_offer": final_offer,
        "active_tab": active_tab,
        "files_by_category": files_by_category,
        "recent_file_categories": dict(file_category_counts),
        "recent_notes": client.notes.order_by("-added_date")[:3],
        "recent_tasks": recent_tasks,
        "recent_activities": client.activities.order_by("-timestamp")[:6],
        "milestones": milestones,
        "task_form": task_form,
        "meeting_form": meeting_form,
        "note_form": note_form,
        "tasks": tasks,
        "meetings": meetings,
        "meetings_rest": meetings_rest,
        "notes": notes,
        "active_type": file_type_label,
        "payments": payments,
        "payment_form": payment_form,
        "offers_with_documents": offers_with_documents,
        "offers": offers
        
    }

    return render(request, "client_card.html", context)






@csrf_exempt
@login_required
def acknowledge_post(request):
    if request.method == "POST":
        post_id = request.POST.get("post_id")
        post = get_object_or_404(InformationPost, id=post_id)

        if request.user not in post.acknowledged_by.all():
            # Dodaj użytkownika do listy potwierdzających
            post.acknowledged_by.add(request.user)
            status = "added"
        
        return JsonResponse({"status": status, "post_id": post.id})
    return JsonResponse({"error": "Invalid request"}, status=400)




@login_required
def dashboard(request):
    user = request.user  # Obiekt użytkownika
    
    # Pobierz posty posortowane od najnowszego do najstarszego
    posts = InformationPost.objects.order_by('-created_at')  # Sortowanie wg daty utworzenia
        # Sprawdź, ile postów wymaga zatwierdzenia
    unacknowledged_posts_count = InformationPost.objects.filter(
        mandatory_to_acknowledge=True
    ).exclude(acknowledged_by=request.user).count()
    # Pobierz najnowszy post (pierwszy w posortowanej liście)
    latest_post = posts.first()
    print("unacknowledged_posts_count",unacknowledged_posts_count)
    # Pobierz pozostałe posty, pomijając najnowszy
    other_posts = posts[1:]
    
    # Paginacja dla pozostałych postów
    paginator = Paginator(other_posts, 6)  # 6+1 postów na stronę
    page_number = request.GET.get('page', 1)
    paginated_posts = paginator.get_page(page_number)
    
    return render(request, 'dashboard.html', {
        'latest_post': latest_post,
        'other_posts': paginated_posts,
        'user': user,
        'unacknowledged_posts_count': unacknowledged_posts_count
    })




@login_required
def post_detail(request, post_id):
    post = get_object_or_404(InformationPost, id=post_id)
    
            # Sprawdź, ile postów wymaga zatwierdzenia
    unacknowledged_posts_count = InformationPost.objects.filter(
        mandatory_to_acknowledge=True
    ).exclude(acknowledged_by=request.user).count()
    
    return render(request, 'post_detail.html', {'post': post, "unacknowledged_posts_count": unacknowledged_posts_count})



@login_required
def clients_view(request):
    if request.user.groups.filter(name="biuro").exists():  # Poprawione sprawdzanie grupy
        leads = Client.objects.filter(status='client')
    else:
        leads = Client.objects.filter(user=request.user, status='client')  # Dopasowanie do pola przypisania użytkownika

    return render(request, 'clients.html', {'leads': leads})


@login_required
def leads_view(request):
    leads = Client.objects.filter(user=request.user, status='lead')
    return render(request, 'leads.html', {'leads': leads})

@login_required
def oferss_view(request):
    leads = Client.objects.filter(user=request.user, status='offer')
    return render(request, 'offers.html', {'leads': leads})



@login_required
def calculator(request):
    # Obsługa query string z lead_id
    lead_id = request.GET.get('lead_id')
    lead = None
    if lead_id:
        try:
            lead = Client.objects.get(id=lead_id, user=request.user)
        except Client.DoesNotExist:
            lead = None
    # Obsługa POST (np. zapis wybranego produktu lub danych)
    if request.method == "POST" and "product_id" in request.POST:
        required_fields = {}
        product = get_object_or_404(Product, id=request.POST["product_id"])
        attributes = product.attributes.all()  # Pobranie powiązanych atrybutów

        context = {
            "selected_product": product,
            "attributes": attributes,
        }
        return render(request, "calculator.html", context)
    programs = SubsidyProgram.objects.all().prefetch_related("options")
    # Dane do wyświetlenia w formularzu
    leads = Client.objects.filter(user=request.user, status='lead')
    categories = ProductCategory.objects.all()  # Wszystkie kategorie produktów

    selected_category = request.GET.get('category')  # ID wybranej kategorii
    selected_product = request.GET.get('product')  # ID wybranego produktu

    # Pobierz produkty dla wybranej kategorii
    products = Product.objects.filter(category_id=selected_category) if selected_category else None

    # Pobierz atrybuty dla wybranego produktu
    attributes = ProductAttribute.objects.filter(product_id=selected_product) if selected_product else None

    form = ClientForm()
    return render(request, 'calculator.html', {
        'leads': leads,
        'categories': categories,
        'products': products,
        'attributes': attributes,
        'form': form,
        'selected_category': selected_category,
        'selected_product': selected_product,
        "programs": programs,
        "selected_lead": lead,  # Przekaż dane wybranego leada do szablonu
    })

def extract_number_from_surface(surface):
    """
    Wyciąga liczbę z ciągu tekstowego powierzchni.
    Obsługuje formaty takie jak: '123m2', '324 m2', '10 metrów', itp.
    """
    if not surface:
        return None
    match = re.search(r'\d+', surface)  # Wyszukaj pierwszą sekwencję cyfr
    if match:
        return Decimal(match.group())  # Zwróć jako Decimal dla dalszych obliczeń
    return None
# /////////////////////////////////////////////////////////////////////////////////////////////////
def calculate_optimal_subsidy(products, max_total_subsidy,program_id):
    """
    Oblicza najkorzystniejszy rozkład dotacji pomiędzy produktami.

    Args:
    products (list): Lista produktów, gdzie każdy produkt to słownik zawierający:
        - name: nazwa produktu
        - price: cena produktu
        - category: kategoria produktu
        - subsidy_percentage: maksymalny procent dotacji
        - max_subsidy_amount: maksymalna kwota dotacji
        - second_subsidy_percentage (opcjonalne): drugi maksymalny procent dotacji
        - second_max_subsidy_amount (opcjonalne): druga maksymalna kwota dotacji
    max_total_subsidy (float): Maksymalna łączna kwota dotacji.

    Returns:
    dict: Rozkład dotacji dla każdego produktu i całkowita dotacja.
    """
    mustPayment = 0
    total_subsidy = 0
    heeter_subsidy = 0
    termo_subsidy = 0
    foto_subsidy = 0
    subsidy_left = 0
    subsidy_distribution = []
    for product in products:
        grossPrice = 0
        # Oblicz maksymalne dotacje dla produktu
        product_price = product["price"]
        grossPrice = (product["vatRate"] * product_price / 100) + product_price
        if Decimal(program_id) == 2:
            product_price = grossPrice
        # Pierwszy próg dotacji
        first_subsidy = min(
            product_price * product["subsidy_percentage"] / 100,
            product["max_subsidy_amount"]
        )
        
        # Drugi próg dotacji, jeśli istnieje
        second_subsidy = 0
        if "second_subsidy_percentage" in product and "second_max_subsidy_amount" in product:
            remaining_price = product_price - (product["max_subsidy_amount"] / product["subsidy_percentage"] * 100)
            if remaining_price > 0:
                second_subsidy = min(
                    remaining_price * product["second_subsidy_percentage"] / 100,
                    product["second_max_subsidy_amount"])
        # Całkowite dofinansowanie dla produktu
        product_subsidy = first_subsidy + second_subsidy
        
        mustPayment = mustPayment + grossPrice
        if "second_max_subsidy_amount" in product:
            if product_subsidy <= (product["max_subsidy_amount"]+product["second_max_subsidy_amount"]):
                subsidy_left = (product["max_subsidy_amount"]+product["second_max_subsidy_amount"]) - product_subsidy
        else:
  
            if product_subsidy <= product["max_subsidy_amount"]:
                subsidy_left = product["max_subsidy_amount"] - product_subsidy

        if len(str(product["category"])) == len("Żródło ogrzewania"):
            heeter_subsidy = subsidy_left
            subsidy_distribution.append({
                "product_category": str(product["category"]),
                "product_name": product["name"],
                "attributes": product["attributes"],
                "product_price": product["price"],
                "subsidy": product_subsidy,
                "subsidy_left": heeter_subsidy,
                "product_id": product["product_id"],
                "product_initial_price": product["product_initial_price"],
                "grossPrice" : grossPrice,
                "vatRate" : product["vatRate"]
            })
        if len(str(product["category"])) == len("Izolacja"):
            if termo_subsidy > 0:
                termo_subsidy = termo_subsidy - product_subsidy
            else:
                termo_subsidy = subsidy_left
            subsidy_distribution.append({
                "product_category": str(product["category"]),
                "product_name": product["name"],
                "attributes": product["attributes"],
                "product_price": product["price"],
                "subsidy": product_subsidy,
                "subsidy_left": termo_subsidy,
                "product_id": product["product_id"],
                "product_initial_price": product["product_initial_price"],
                "grossPrice" : grossPrice,
                "vatRate" : product["vatRate"]
            })
        if len(str(product["category"])) == len("Fotowoltaika i magazyn energii"):
            foto_subsidy = subsidy_left
            subsidy_distribution.append({
                "product_category": str(product["category"]),
                "product_name": product["name"],
                "attributes": product["attributes"],
                "product_price": product["price"],
                "subsidy": product_subsidy,
                "subsidy_left": foto_subsidy,
                "product_id": product["product_id"],
                "product_initial_price": product["product_initial_price"],
                "grossPrice" : grossPrice,
                "vatRate" : product["vatRate"]
            })
        total_subsidy += product_subsidy
        
    # Ograniczenie całkowitej dotacji
    if total_subsidy > max_total_subsidy:
        scaling_factor = max_total_subsidy / total_subsidy
        for item in subsidy_distribution:
            item["subsidy"] = round(item["subsidy"] * scaling_factor, 2)
        total_subsidy = max_total_subsidy

    return {
            "subsidy_distribution": subsidy_distribution,
            "total_subsidy": round(total_subsidy, 2),
            "termo_subsidy": round(termo_subsidy, 2),
            "heeter_subsidy": round(heeter_subsidy, 2),
            "mustPayment": mustPayment
        }
    


def calculate_maximum_subsidy(products, max_total_subsidy,program_id):
    mustPayment = 0
    subsidy_distribution = []
    total_subsidy = 0
    heeter_subsidy = 0
    termo_subsidy = 0
    foto_subsidy = 0
    heeter_product = 0
    termo_product = 0
    foto_product = 0
    lock = 0
    lockFoto = 0
    maglock = 0
    subsidy_total_left = max_total_subsidy
    print(subsidy_total_left)
    for product in products:
        grossPrice = 0
        if len(str(product["category"])) == len("Żródło ogrzewania"):
            heeter_product += 1
        if len(str(product["category"])) == len("Izolacja"):
            termo_product += 1
        if len(str(product["category"])) == len("Fotowoltaika i magazyn energii"):
            if product["max_subsidy_amount"] != 0:
                foto_subsidy = product["max_subsidy_amount"]
            print(foto_subsidy)
            foto_product += 1
    
    for product in products:
        if lock < 1 and heeter_product > 0 and len(str(product["category"])) == len("Żródło ogrzewania"):
            lock += 1
            product_price = product["max_subsidy_amount"] / product["subsidy_percentage"] * 100
            product_subsidy = product["max_subsidy_amount"]
            if "second_subsidy_percentage" in product and "second_max_subsidy_amount" in product:
                product_price = product_price + (product["second_max_subsidy_amount"] / product["second_subsidy_percentage"] * 100)
                if product_price < float(product["product_initial_price"]):
                    product_price = float(product["product_initial_price"])
                product_subsidy = product_subsidy + product["second_max_subsidy_amount"]
            subsidy_total_left = subsidy_total_left - product_subsidy
            total_subsidy += product_subsidy
            heeter_product += product_subsidy
            if product_price < float(product["product_initial_price"]):
                product_price = float(product["product_initial_price"])
            grossPrice = (product["vatRate"] * product_price / 100) + product_price
            mustPayment = mustPayment + grossPrice
            subsidy_distribution.append({
                "product_category": str(product["category"]),
                "product_name": product["name"],
                "attributes": product["attributes"],
                "product_price": product_price,
                "subsidy": product_subsidy,
                "subsidy_left": 0,
                "maxSub" : True,
                "product_id": product["product_id"],
                "product_initial_price": product["product_initial_price"],
                "grossPrice" : grossPrice,
                "vatRate" : product["vatRate"]
            })
    if termo_product > 0:
        subsidy_termo_each_left = subsidy_total_left / termo_product
        if foto_product > 0:
            subsidy_termo_each_left = (subsidy_total_left - foto_subsidy / termo_product)
    for product in products:
        if termo_product > 0 and len(str(product["category"])) == len("Izolacja"):
            
            if product["max_subsidy_amount"] > subsidy_termo_each_left:
                product_price = subsidy_termo_each_left / product["subsidy_percentage"] * 100    
                product_subsidy = subsidy_termo_each_left

            else:
                product_price = product["max_subsidy_amount"] / product["subsidy_percentage"] * 100
                product_subsidy = product["max_subsidy_amount"]
            if "second_subsidy_percentage" in product and "second_max_subsidy_amount" in product:
                if (product["max_subsidy_amount"] + product["second_max_subsidy_amount"]) > subsidy_termo_each_left:
                    product_price = subsidy_termo_each_left / product["subsidy_percentage"] * 100    
                    product_subsidy = subsidy_termo_each_left
                else:
                    product_price = product_price + (product["second_max_subsidy_amount"] / product["second_subsidy_percentage"] * 100)
                    product_subsidy = product_subsidy + product["second_max_subsidy_amount"]
            subsidy_total_left = subsidy_total_left - product_subsidy
            total_subsidy += product_subsidy
            termo_product += product_subsidy
            if product_price < float(product["product_initial_price"]):
                product_price = float(product["product_initial_price"])
            grossPrice = (product["vatRate"] * product_price / 100) + product_price
            mustPayment = mustPayment + grossPrice
            subsidy_distribution.append({
                "product_category": str(product["category"]),
                "product_name": product["name"],
                "attributes": product["attributes"],
                "product_price": product_price,
                "subsidy": product_subsidy,
                "subsidy_left": 0,
                "maxSub" : True,
                "product_id": product["product_id"],
                "product_initial_price": product["product_initial_price"],
                "grossPrice" : grossPrice,
                "vatRate" : product["vatRate"]
            })
        if lockFoto <= 1 and foto_product > 0 and len(str(product["category"])) == len("Fotowoltaika i magazyn energii"):
            if str(product["name"]) in ["Magazyn Energii"]:
                maglock += 1
            print(program_id, str(product["name"]) not in ["Magazyn Energii"])
            if Decimal(program_id) != 2 or str(product["name"]) not in ["Magazyn Energii"]:
                lockFoto += 1
            product_price = product["max_subsidy_amount"] / product["subsidy_percentage"] * 100
            if Decimal(program_id) == 2:
                product_price = product_price / (1 + product["vatRate"] / 100)
            if product_price < float(product["product_initial_price"]):
                product_price = float(product["product_initial_price"])
            product_subsidy = product["max_subsidy_amount"]
            if "second_subsidy_percentage" in product and "second_max_subsidy_amount" in product:
                product_price = product_price + (product["second_max_subsidy_amount"] / product["second_subsidy_percentage"] * 100)
                if product_price < product["price"]:
                    product_price = product["price"]
                product_subsidy = product_subsidy + product["second_max_subsidy_amount"]
            if str(product["name"]) in ["Magazyn Energii"] and Decimal(program_id) != 2:
                product_subsidy = 0
            if maglock > 1: 
                product_subsidy = 0
            if heeter_product == 0 and Decimal(program_id) != 2:
                product_subsidy = 0
            subsidy_total_left = subsidy_total_left - product_subsidy
            total_subsidy += product_subsidy
            foto_product += product_subsidy
            grossPrice = (product["vatRate"] * product_price / 100) + product_price
            mustPayment = mustPayment + grossPrice
            subsidy_distribution.append({
                "product_category": str(product["category"]),
                "product_name": product["name"],
                "attributes": product["attributes"],
                "product_price": product_price,
                "subsidy": product_subsidy,
                "subsidy_left": 0,
                "maxSub" : True,
                "product_id": product["product_id"],
                "product_initial_price": product["product_initial_price"],
                "grossPrice" : grossPrice,
                "vatRate" : product["vatRate"]
            })
        elif len(str(product["category"])) == len("Fotowoltaika i magazyn energii"):
            print("leci")
            product_subsidy = 0
            total_subsidy += product_subsidy
            grossPrice = (product["vatRate"] * product["price"] / 100) + product["price"]
            mustPayment = mustPayment + grossPrice
            subsidy_distribution.append({
                "product_category": str(product["category"]),
                "product_name": product["name"],
                "attributes": product["attributes"],
                "product_price": product["price"],
                "subsidy": product_subsidy,
                "subsidy_left": 0,
                "maxSub" : True,
                "product_id": product["product_id"],
                "product_initial_price": product["product_initial_price"],
                "grossPrice" : grossPrice,
                "vatRate" : product["vatRate"]
            })

    return {
            "subsidy_distribution": subsidy_distribution,
            "total_subsidy": round(total_subsidy, 2),
            "termo_subsidy": round(termo_subsidy, 2),
            "heeter_subsidy": round(heeter_subsidy, 2),
            "mustPayment": mustPayment
        }
    
@login_required
def complete_task(request, task_id):
    task = get_object_or_404(Task, id=task_id, client__user=request.user)
    task.completed = True
    task.save()
    messages.success(request, "Zadanie oznaczone jako wykonane.")
        # Pobierz aktywną zakładkę z parametru URL (lub ustaw domyślną, jeśli nie istnieje)
    active_tab = request.GET.get('tab', 'overview')
    create_notification("oznaczono zadanie jako 'wykonane'", request.user, client_id=task.client.id)
        # Przekieruj do widoku karty klienta z aktywną zakładką
    return redirect(f"{reverse('client_card', args=[task.client.id])}?tab={active_tab}")

@login_required
def mark_meeting_occurred(request, meeting_id):
    meeting = get_object_or_404(Meeting, id=meeting_id, client__user=request.user)
    occurred = request.POST.get("occurred")
    if occurred == "true":
        occurred = True
    else:
        occurred= False
    print(occurred)
    if request.method == "POST":
        note = request.POST.get("note", "").strip()
        if not note:
            messages.error(request, "Notatka jest wymagana, aby oznaczyć spotkanie jako odbyte.")
            return redirect("client_meetings", client_id=meeting.client.id)

        meeting.occurred = occurred
        meeting.note = note
        meeting.save()
        messages.success(request, "Spotkanie oznaczone jako odbyte.")
    
    active_tab = request.GET.get('tab', 'overview')
    print(active_tab)
        # Przekieruj do widoku karty klienta z aktywną zakładką
    return redirect(f"{reverse('client_card', args=[meeting.client.id])}?tab={active_tab}")


def extract_placeholders(doc_path):
    """Funkcja do skanowania pliku DOCX i wyszukiwania placeholderów w paragrafach i tabelach"""
    doc = docx.Document(doc_path)
    placeholders = set()
    pattern = re.compile(r"\{\{\s*(.*?)\s*\}\}")

    # Skanowanie paragrafów w dokumencie
    for para in doc.paragraphs:
        matches = pattern.findall(para.text)
        placeholders.update(matches)

    # Skanowanie tabel w dokumencie
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = pattern.findall(cell.text)
                placeholders.update(matches)

    return sorted(placeholders)  # Zwracamy posortowaną listę dla lepszej czytelności

def upload_template(request):
    templates = DocumentTemplate.objects.all()

    if request.method == "POST":
        form = DocumentTemplateForm(request.POST, request.FILES)
        
        if form.is_valid():
            template = form.save(commit=False)  # Nie zapisujemy od razu
            
            # **WAŻNE**: Musimy zapisać plik, aby móc go odczytać
            template.save()
            template.refresh_from_db()  # Upewniamy się, że Django widzi plik

            # Pobieramy rzeczywistą ścieżkę do pliku
            file_path = os.path.join(settings.MEDIA_ROOT, template.file.name)

            # Sprawdzamy, czy plik rzeczywiście istnieje
            if not os.path.exists(file_path):
                messages.error(request, "Błąd: Plik nie został poprawnie zapisany.")
                return redirect(request.path)

            # Teraz możemy bezpiecznie odczytać plik i wyodrębnić placeholdery
            placeholders = extract_placeholders(file_path)
            template.placeholders = placeholders
            template.save()

            messages.success(request, "Szablon został dodany pomyślnie!")
            return redirect(request.path)  # Przekierowanie po dodaniu

    else:
        form = DocumentTemplateForm()

    return render(request, "upload_template.html", {"form": form, "templates": templates})
    
@login_required
def delete_template(request, template_id):
    template = get_object_or_404(DocumentTemplate, id=template_id)
    
    if request.method == "POST":
        template.delete()
        messages.success(request, "Szablon został usunięty pomyślnie!")
        return redirect("upload_template")  # Jeśli w urls.py masz name='upload_template'

    messages.error(request, "Nieprawidłowe żądanie usunięcia!")
    return redirect("upload_template")  # Jeśli w urls.py masz name='upload_template'


def przypisz_wojewodztwo(kod_pocztowy):
    wojewodztwa = {
        ("00", "01", "02", "03", "04", "05", "06", "07"): "mazowieckie",
        ("08", "09"): "mazowieckie",  # Warszawa i okolice

        ("10", "11", "12", "13", "14"): "warmińsko-mazurskie",

        ("15", "16", "17", "18", "19"): "podlaskie",

        ("20", "21", "22"): "lubelskie",

        ("23", "24"): "lubelskie",
        ("25", "26"): "świętokrzyskie",

        ("27", "28", "29"): "świętokrzyskie",

        ("30", "31", "32", "33", "34"): "małopolskie",

        ("35", "36", "37", "38", "39"): "podkarpackie",

        ("40", "41", "42", "43", "44"): "śląskie",

        ("45", "46", "47", "48", "49"): "opolskie",

        ("50", "51", "52", "53", "54", "55"): "dolnośląskie",
        ("56", "57", "58", "59"): "dolnośląskie",

        ("60", "61", "62", "63", "64"): "wielkopolskie",
        ("65", "66"): "wielkopolskie",

        ("67", "68"): "lubuskie",
        ("69",): "lubuskie",

        ("70", "71", "72", "73", "74"): "zachodniopomorskie",
        ("75", "76"): "zachodniopomorskie",

        ("77", "78"): "pomorskie",
        ("79", "80", "81", "82", "83", "84"): "pomorskie",

        ("85", "86", "87", "88", "89"): "kujawsko-pomorskie",

        ("90", "91", "92", "93", "94", "95", "96", "97", "98", "99"): "łódzkie"
    }

    try:
        pierwsze_dwie_cyfry = kod_pocztowy[:2]
        for klucze, wojewodztwo in wojewodztwa.items():
            if pierwsze_dwie_cyfry in klucze:
                return wojewodztwo
        return "Nieznane województwo"
    except IndexError:
        return "Niepoprawny kod pocztowy"
    
@login_required
@csrf_exempt
@login_required
@csrf_exempt
def save_final_offer(request):
    if request.method == "POST":
        client_id = request.POST.get("client_id")
        offer_id = request.POST.get("selected_offer")

        if not client_id or not offer_id:
            return JsonResponse({"success": False, "error": "Brak wymaganych danych."})

        # Pobranie klienta i oferty
        client = get_object_or_404(Client, id=client_id)
        offer = get_object_or_404(Offer, id=offer_id)

        # Przypisanie finalnej oferty do klienta
        client.final_offer = offer
        client.contract_value = offer.total_price
        client.margin = offer.total_margin  # Pobranie marży z oferty
        client.status = "client"  # Zmiana statusu klienta na "Klient"
        client.milestone = "agreement"
        client.save()

        return JsonResponse({"success": True})





def get_event_to_calendar(request):
    # Pobranie wszystkich zadań
    tasks = Task.objects.select_related("client").all().values(
        "id", "due_date", "text", "author__username", "added_date", "completed",
        "client__id", "client__first_name", "client__last_name", "completed_at"
    )

    # Pobranie wszystkich spotkań, dołączając dane klienta
    meetings = Meeting.objects.select_related("client").all().values(
        "id", "meeting_date", "description", "note", "author__username", "occurred",
        "client__id", "client__first_name", "client__last_name","note_added_at"
    )

    # Formatowanie zadań
    formatted_tasks = [
        {
            "id": task["id"],  # Dodanie ID zadania
            "date": timezone.localtime(task["due_date"]).strftime("%Y-%m-%d"),
            "time": timezone.localtime(task["due_date"]).strftime("%H:%M"),  # Dodajemy godzinę
            "title": task["text"] if task["text"] else "Brak opisu",
            "author": task["author__username"] if task["author__username"] else "Nieznany autor",
            "added_date": timezone.localtime(task["added_date"]).strftime("%Y-%m-%d %H:%M"),
            "completed": task["completed"],
            "completed_at": task["completed_at"],
            "client_name": f"{task['client__first_name']} {task['client__last_name']}",
            "client_link": f"/klienci/klient/{task['client__id']}/"
        }
        for task in tasks if task["due_date"] is not None
    ]

    # Formatowanie spotkań
    formatted_meetings = [
        {
            "id": meeting["id"],  # Dodanie ID spotkania
            "date": timezone.localtime(meeting["meeting_date"]).strftime("%Y-%m-%d"),
             "time": timezone.localtime(meeting["meeting_date"]).strftime("%H:%M"),  # Dodajemy godzinę
            "title":  meeting["description"],
            "author": meeting["author__username"] if meeting["author__username"] else "Nieznany autor",
            "note": meeting["note"],
            "note_added_at": meeting["note_added_at"],
            "occurred": meeting["occurred"],
            "client_name": f"{meeting['client__first_name']} {meeting['client__last_name']}",
            "client_link": f"/klienci/klient/{meeting['client__id']}/"
        }
        for meeting in meetings if meeting["meeting_date"] is not None
    ]

    # Zwrot danych w formacie JSON
    return JsonResponse({"tasks": formatted_tasks, "meetings": formatted_meetings})



@csrf_exempt
def mark_task_completed(request, task_id):
    if request.method == "POST":
        try:
            task = Task.objects.get(id=task_id)
            task.completed = True
            task.completed_at = timezone.now()
            task.save()
            # Powiadamiamy tylko odpowiedniego użytkownika
            print("jestem")
            create_notification("oznaczono zadanie jako 'wykonane'", request.user, client_id=task.client.id)
            return JsonResponse({"success": True})
        except Task.DoesNotExist:
            return JsonResponse({"success": False, "error": "Zadanie nie istnieje."})
    return JsonResponse({"success": False, "error": "Nieprawidłowa metoda żądania."})



@csrf_exempt
def update_meeting_status(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            meeting_id = data.get("meeting_id")
            occurred = data.get("occurred")
            note = data.get("note", "").strip()

            if not meeting_id or note == "":
                return JsonResponse({"success": False, "error": "Brak ID spotkania lub notatki"}, status=400)

            meeting = Meeting.objects.get(id=meeting_id)
            meeting.occurred = occurred
            meeting.note = note
            meeting.save()

            return JsonResponse({"success": True})
        except Exception as e:
            return JsonResponse({"success": False, "error": str(e)}, status=500)

    return JsonResponse({"success": False, "error": "Nieprawidłowa metoda"}, status=405)


@login_required
def notifications_list(request):
    """ Wyświetla wszystkie powiadomienia użytkownika z paginacją """
    notifications = Notification.objects.filter(recipient=request.user).order_by('-created_at')

    # Paginacja – 20 powiadomień na stronę
    paginator = Paginator(notifications, 20)  
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    return render(request, 'notifications_list.html', {'page_obj': page_obj})

@login_required
def get_recent_notifications(request):
    notifications = Notification.objects.filter(recipient=request.user).order_by('-created_at')[:10]
    data = [
        {"title": "🔔 " + n.title, "message": n.message, "created_at": n.created_at, "is_read": n.is_read, "id" :n.id,}
        for n in notifications
    ]
    return JsonResponse({"notifications": data})

@csrf_exempt
def mark_notification_as_read(request, notification_id):
    """Oznacza powiadomienie jako przeczytane"""
    if request.method == "POST":
        try:
            notification = Notification.objects.get(id=notification_id)
            notification.is_read = True
            notification.save()
            return JsonResponse({"success": True, "message": "Powiadomienie oznaczone jako przeczytane"})
        except Notification.DoesNotExist:
            return JsonResponse({"success": False, "error": "Powiadomienie nie istnieje"}, status=404)
    return JsonResponse({"success": False, "error": "Nieprawidłowa metoda żądania"}, status=400)




def create_notification(action, creator, client_id):
    """
    Tworzy powiadomienie dla odpowiedniego użytkownika w zależności od tego, kto dodaje dane.

    - Jeśli `creator` jest handlowcem, powiadomienie dostaje biuro.
    - Jeśli `creator` jest biurem, powiadomienie dostaje przypisany handlowiec.
    
    :param action: string (np. "zadanie", "notatka", "płatność")
    :param creator: User (obiekt użytkownika, który tworzy)
    :param client_id: int (ID klienta, jeśli dotyczy konkretnego klienta)
    """


    biuro_users = User.objects.filter(groups__name="biuro")

    recipients = []

    if creator.groups.filter(name="handlowcy").exists():
        client = Client.objects.get(id=client_id)
        recipients = list(biuro_users)
        message = f"Handlowiec {creator.username} dodał {action} dla klienta {client.first_name} {client.last_name}."


    else:
        if client_id:
            try:
                client = Client.objects.get(id=client_id)
                if client.user:
                    recipients = [client.user]

                    message = f"Biuro dodało {action} dla klienta {client.first_name} {client.last_name}."

                else:

                    return  
            except Client.DoesNotExist:

                return  

    if not recipients:

        return  # Nie tworzymy powiadomienia, jeśli nie ma odbiorcy
    print("📌 Lista odbiorców:", recipients)
    for recipient in recipients:
        if recipient is None:
            continue  

        Notification.objects.create(
            user=creator,
            title= (f"{action}").upper(),
            message=message,
            recipient=recipient  # To musi być poprawne!
        )
        
def get_coordinates_from_postal_code(postal_code, country="PL"):
    """
    Pobiera szerokość i długość geograficzną na podstawie kodu pocztowego.
    """
    url = f"https://nominatim.openstreetmap.org/search?postalcode={postal_code}&country={country}&format=json"
    headers = {"User-Agent": "MyPythonApp/1.0 (contact@example.com)"}


    try:
        response = requests.get(url, headers=headers)  # Dodajemy User-Agent
        response.raise_for_status()
        data = response.json()

        if data:
            lat = float(data[0]['lat'])
            lon = float(data[0]['lon'])
            return lat, lon
        else:
            raise ValueError("Nie znaleziono współrzędnych dla podanego kodu pocztowego.")

    except requests.exceptions.RequestException as e:
        print(f"❌ Błąd połączenia z OpenStreetMap: {e}")
        return None, None
    except (ValueError, IndexError, KeyError) as e:
        print(f"⚠️ Błąd przetwarzania danych OpenStreetMap: {e}")
        return None, None



def get_solar_radiation(lat, lon):
    """
    Pobiera roczne nasłonecznienie w kWh/m² na podstawie współrzędnych.
    """
    url = f"https://power.larc.nasa.gov/api/temporal/climatology/point?parameters=ALLSKY_SFC_SW_DWN&community=RE&longitude={lon}&latitude={lat}&format=JSON"

    try:
        response = requests.get(url)
        response.raise_for_status()
        data = response.json()

        if 'properties' in data and 'parameter' in data['properties']:
            solar_data = data['properties']['parameter'].get('ALLSKY_SFC_SW_DWN', {})

            # Liczba dni w każdym miesiącu
            days_in_month = {
                "JAN": 31, "FEB": 28, "MAR": 31, "APR": 30, "MAY": 31, "JUN": 30,
                "JUL": 31, "AUG": 31, "SEP": 30, "OCT": 31, "NOV": 30, "DEC": 31
            }

            # Obliczenie rocznego nasłonecznienia
            yearly_radiation = sum(float(solar_data[month]) * days for month, days in days_in_month.items())

            return round(yearly_radiation, 2)  # Zaokrąglenie do 2 miejsc po przecinku
        else:
            raise ValueError("Nie udało się pobrać danych o nasłonecznieniu.")

    except requests.exceptions.RequestException as e:
        print(f"❌ Błąd połączenia z NASA POWER API: {e}")
        return None
    except (ValueError, KeyError, TypeError) as e:
        print(f"⚠️ Błąd przetwarzania danych NASA POWER API: {e}")
        return None



def get_solar_data(postal_code):
    """
    Pobiera szerokość i długość geograficzną oraz nasłonecznienie roczne na podstawie kodu pocztowego.
    
    :param postal_code: Kod pocztowy (str)
    :return: Tuple (lat, lon, annual_solar_radiation) lub (None, None, None) w przypadku błędu
    """
    try:
        # Pobranie współrzędnych geograficznych
        lat, lon = get_coordinates_from_postal_code(postal_code)
        
        if lat is None or lon is None:
            print(f"⚠️ Nie udało się uzyskać współrzędnych dla kodu pocztowego: {postal_code}")
            return None, None, None
        
        # Pobranie średniego dziennego nasłonecznienia z NASA POWER API
        daily_radiation = get_solar_radiation(lat, lon)
        
        if daily_radiation is not None:
            # Przekształcenie na roczne nasłonecznienie
            annual_radiation = round(daily_radiation, 2)
            

            return lat, lon, annual_radiation
        else:
            print(f"⚠️ Nie udało się obliczyć nasłonecznienia dla {postal_code}")
            return lat, lon, None
    
    except Exception as e:
        print(f"🚨 Nieoczekiwany błąd: {e}")
        return None, None, None

def client_registration(request):
    if request.method == "POST":
        form = PreleadForm(request.POST)
        if form.is_valid():
            client = form.save(commit=False)
            client.save()
            messages.success(request, "Twoje zgłoszenie zostało zapisane. Skontaktujemy się wkrótce!")
            return JsonResponse({"success": True, "message": "Twoje zgłoszenie zostało zapisane. Skontaktujemy się wkrótce!"})
        else:
            return JsonResponse({"success": False, "errors": form.errors}, status=400)
    else:
        form = PreleadForm()

    return render(request, "client_registration.html", {"form": form})




def parcel_points(request):
    data = []
    for p in Parcel.objects.filter(latitude__isnull=False, longitude__isnull=False):
        data.append({
            "latitude": p.latitude,
            "longitude": p.longitude,
            "town": p.town,
            "plot_number": p.plot_number,
            "area": p.area,
            "status": p.lead.status if p.lead else None,
            "first_name": p.lead.first_name if p.lead else "",
            "last_name": p.lead.last_name if p.lead else "",
            "lead_id": p.lead.id if p.lead else None
        })
    return JsonResponse(data, safe=False)

def geocode_parcels_view(request):
    if not request.user.is_staff:
        return redirect('/')  # lub zwróć 403

    updated = 0
    skipped = 0

    parcels = Parcel.objects.filter(latitude__isnull=True, longitude__isnull=True)

    for parcel in parcels:
        lat, lon = get_coordinates(
            voivodeship=parcel.voivodeship,
            county=parcel.county,
            town=parcel.town,
            precinct=parcel.precinct,
            plot_number=parcel.plot_number
        )

        if lat and lon:
            parcel.latitude = lat
            parcel.longitude = lon
            parcel.save()
            updated += 1
        else:
            skipped += 1

        time.sleep(1)  # rate limit

    messages.success(request, f"Zaktualizowano {updated} działek, pominięto {skipped}.")
    return redirect(request.META.get("HTTP_REFERER", "/"))
